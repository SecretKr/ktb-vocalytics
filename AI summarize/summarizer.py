import csv
import glob
import os
import re

import docx
import ollama
from fuzzywuzzy import (  # Imported for add_keyword_highlighting, though not directly for fuzzy correction here
    fuzz, process)


# Function to read text from .docx file
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to append text to an existing .docx file
def append_to_docx(file_path, content, model_name):
    doc = docx.Document(file_path) # Open existing document
    doc.add_paragraph(f"\n--- สรุปเนื้อหาโดย AI (เพิ่มเติม) โดยใช้โมเดล {model_name} ---\n") # Add a separator
    doc.add_paragraph(content)
    doc.save(file_path) # Save changes to the same file

# Function to read keywords from a CSV file
def read_keywords_from_csv(csv_file_path):
    keywords = []
    try:
        with open(csv_file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if row: # Ensure row is not empty
                    keywords.extend([kw.strip() for kw in row if kw.strip()])
    except FileNotFoundError:
        print(f"Keyword file not found: {csv_file_path}")
    return list(set(keywords)) # Return unique keywords

# Function to add keyword highlighting with fuzzy matching
def add_keyword_highlighting(text, keywords, threshold=80):
    highlighted_text = text
    words = re.findall(r'\b\w+\b', text) # Find all words in the text
    
    for keyword in keywords:
        for word in words:
            if fuzz.ratio(word.lower(), keyword.lower()) >= threshold:
                # Replace the matched word with its highlighted version
                highlighted_text = re.sub(r'\b' + re.escape(word) + r'\b', r'**' + word + r'**', highlighted_text, flags=re.IGNORECASE)
    return highlighted_text

if __name__ == '__main__':
    script_dir = os.path.dirname(__file__)
    input_directory = os.path.join(script_dir, "transcript_with_highlight")
    keywords_directory = os.path.join(script_dir, os.pardir, "keywords")
    
    os.makedirs(input_directory, exist_ok=True)
    
    print(f"Attempting to find .docx files in: {os.path.abspath(input_directory)}")
    # Exclude temporary Word files starting with "~$"
    all_docx_files = glob.glob(os.path.join(input_directory, "*.docx"))
    docx_files = [f for f in all_docx_files if not os.path.basename(f).startswith('~$')]
    print(f"Files found: {docx_files}")

    if not docx_files:
        print(f"No .docx files found in {os.path.abspath(input_directory)}")
    else:
        # Create an Ollama client pointing to the Docker container
        client = ollama.Client(host='http://localhost:11434')

        # Define the model to use (changed to a potentially smaller model)
        ollama_model = 'gemma3:4b'

        # Load all keywords from both files
        personal_loan_keywords = read_keywords_from_csv(os.path.join(keywords_directory, "personal_loan.csv"))
        debit_card_keywords = read_keywords_from_csv(os.path.join(keywords_directory, "debit_card.csv"))
        keywords_to_use = list(set(personal_loan_keywords + debit_card_keywords))

        for file_path in docx_files:
            file_name = os.path.basename(file_path)
            print(f"Processing {file_name} for summarization using model: {ollama_model}...")
            
            try:
                document_text = read_docx(file_path)

                if not document_text.strip():
                    print(f"The document {file_name} is empty. Skipping summarization.")
                    continue
                else:
                    # The improved prompt for the AI model
                    prompt_template = """

                    กรุณาสรุปเนื้อหาทั้งหมดเป็นภาษาไทยเท่านั้น

                    คุณคือผู้เชี่ยวชาญด้านการวิเคราะห์บทสนทนา Call Center สำหรับงานขายสินเชื่อและผลิตภัณฑ์ทางการเงิน
                    งานของคุณคือการสรุปและวิเคราะห์บทสนทนาอย่างเป็นระบบ โดยพิจารณาจากข้อความที่ได้จากระบบ Voice to Text (แปลงจากเสียงพนักงานคุยกับลูกค้า)
                    กรุณาสรุปเนื้อหาจากบทสนทนาอย่างละเอียด โดยแสดงผลในรูปแบบหัวข้อดังนี้:

                    1. การทักทายและเริ่มต้นบทสนทนา
                    - ระบุว่ามีการกล่าวคำทักทายหรือไม่ เช่น "สวัสดีค่ะ/ครับ", "ขออนุญาตเรียนสาย", "รบกวนเวลาสักครู่", "สะดวกคุยไหม"

                    2. รายละเอียดของผลิตภัณฑ์ที่เสนอ
                    - พนักงานนำเสนอสินค้าหรือบริการอะไร เช่น สินเชื่อส่วนบุคคล, บัตรเครดิต ฯลฯ
                    - มีการอธิบายจุดเด่น เงื่อนไข หรือประโยชน์ของผลิตภัณฑ์หรือไม่

                    3. การแสดงพฤติกรรม Responsible Lending (ถ้ามี)
                    - มีการกล่าวคำเตือนหรือกระตุ้นจริยธรรมในการกู้ เช่น "ควรกู้เท่าที่จำเป็น", "วางแผนการใช้เงินก่อนตัดสินใจ", ฯลฯ หรือไม่

                    4. การกล่าวขอบคุณและการปิดท้ายบทสนทนา
                    - มีการกล่าวคำ “ขอบคุณ” หรือสื่อสารปิดท้ายอย่างสุภาพหรือไม่

                    5. สรุปใจความสำคัญของบทสนทนา
                    - สรุปภาพรวมของสิ่งที่พนักงานและลูกค้าพูดคุยหรือเสนอขาย
                    - ระบุว่าลูกค้าแสดงความสนใจหรือปฏิเสธ
                    - ระบุลำดับเหตุการณ์สั้น ๆ อย่างชัดเจน

                    ข้อควรปฏิบัติ:
                    - สรุปจากเนื้อหาในข้อความเท่านั้น ห้ามเดาหรือเสริมสิ่งที่ไม่มีอยู่จริง
                    - หากมีคำพูดขาดหายหรือฟังไม่ชัด ให้ใส่ข้อความว่า [ฟังไม่ชัด] หรือ [ขาดหาย]
                    - หากใช้ภาษาที่ชัดเจน เป็นกลาง และเหมาะสำหรับใช้ในรายงานทางธุรกิจมีคำพูดขาดหายหรือฟังไม่ชัด ให้ใส่ข้อความว่า [ฟังไม่ชัด] หรือ [ขาดหาย]

                    **โปรดสรุปเนื้อหาทั้งหมดเป็นภาษาไทยเท่านั้น**

                    {document_text}
                    """
                    prompt = prompt_template.format(document_text=document_text)
                    print(f"Sending text from {file_name} to AI for summarization...")
                    
                    # Collect stream content
                    summarized_content = ""
                    stream = client.chat(
                        model=ollama_model, # Use the defined model variable
                        messages=[{'role': 'user', 'content': prompt}],
                        stream=True,
                    )
                    for chunk in stream:
                        summarized_content += chunk['message']['content']
                    
                    print(f"\n--- สรุปเนื้อหาโดย AI สำหรับ {file_name} (Model: {ollama_model}) ---\n")
                    print(summarized_content.strip())
                    print("\n" + "-"*30)

                    # Add keyword highlighting
                    highlighted_content = add_keyword_highlighting(summarized_content, keywords_to_use)

                    # Append summarized and highlighted content to the original file
                    append_to_docx(file_path, highlighted_content, ollama_model)
                    print(f"Summarized and highlighted content appended to {file_path}")

            except FileNotFoundError:
                print(f"Error: File not found at {file_path}")
            except Exception as e:
                print(f"An error occurred while processing {file_name}: {e}")
