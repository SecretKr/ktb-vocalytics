import csv
import re
import os
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# === Config ===
TRANSCRIPT_FILE = 'transcript/transcript.csv'
KEYWORDS_FILE = 'keywords/personal_loan.csv'

DOC_TITLE = 'Transcript with Highlights'
OUTPUT_FILE = 'transcript_with_highlights.docx'

# === Step 1: Load and flatten transcript ===
def load_transcript(file_path):
    with open(file_path, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        chunks = [row['text'] for row in reader]
        full_text = ' '.join(chunks)
    return ' ' + full_text

# === Step 2: Load keyword sequences ===
def load_keyword_patterns(file_path):
    keyword_groups = {}
    with open(file_path, newline='', encoding='utf-8') as f:
        reader = csv.reader(f)
        header = next(reader)
        for row in reader:
            if len(row) < 13:
                continue
            group_name = row[1].strip()
            color_hex = row[12].strip()
            keywords = [w.strip() for w in row[3:7] if w.strip()]

            if not group_name or not keywords:
                continue

            if group_name not in keyword_groups:
                keyword_groups[group_name] = {
                    "patterns": [],
                    "found_words": []
                }
            
            keyword_string = " ".join(keywords)
            pattern = r'.*?'.join(map(re.escape, keywords))
            try:
                rgb = hex_to_rgb_tuple(color_hex)
            except ValueError as e:
                print(f"⚠️ Skipping invalid color for group '{group_name}': {color_hex} ({e})")
                continue
            
            keyword_groups[group_name]["patterns"].append((keyword_string, re.compile(pattern, re.IGNORECASE | re.DOTALL), rgb))
            
    return keyword_groups

def hex_to_rgb_tuple(hex_color):
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        raise ValueError(f"Invalid hex color: {hex_color}")
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return (r, g, b)

def get_closest_wd_color_index(rgb_tuple):
    color_map = {
        (0, 0, 0): WD_COLOR_INDEX.BLACK,
        (0, 0, 255): WD_COLOR_INDEX.BLUE,
        (0, 255, 0): WD_COLOR_INDEX.BRIGHT_GREEN,
        (0, 0, 128): WD_COLOR_INDEX.DARK_BLUE,
        (128, 0, 0): WD_COLOR_INDEX.DARK_RED,
        (128, 128, 0): WD_COLOR_INDEX.DARK_YELLOW,
        (192, 192, 192): WD_COLOR_INDEX.GRAY_25,
        (128, 128, 128): WD_COLOR_INDEX.GRAY_50,
        (0, 128, 0): WD_COLOR_INDEX.GREEN,
        (255, 0, 255): WD_COLOR_INDEX.PINK,
        (255, 0, 0): WD_COLOR_INDEX.RED,
        (0, 128, 128): WD_COLOR_INDEX.TEAL,
        (64, 224, 208): WD_COLOR_INDEX.TURQUOISE, # Approximate
        (128, 0, 128): WD_COLOR_INDEX.VIOLET,
        (255, 255, 255): WD_COLOR_INDEX.WHITE,
        (255, 255, 0): WD_COLOR_INDEX.YELLOW,
    }

    min_distance = float('inf')
    closest_wd_color = WD_COLOR_INDEX.YELLOW # Default if no close match

    for wd_rgb, wd_index in color_map.items():
        distance = sum([(a - b) ** 2 for a, b in zip(rgb_tuple, wd_rgb)]) ** 0.5
        if distance < min_distance:
            min_distance = distance
            closest_wd_color = wd_index
    return closest_wd_color

# === Step 3: Create DOCX with highlights ===
def create_docx_and_highlight(full_text, keyword_groups):
    document = Document()
    document.add_heading(DOC_TITLE, 0)
    
    # Add the full text to a paragraph
    p = document.add_paragraph()
    p.add_run(full_text)

    # Highlight the text
    for group_name, data in keyword_groups.items():
        for keyword_string, pattern, rgb_color in data["patterns"]:
            for match in pattern.finditer(full_text):
                start, end = match.start(), match.end()
                if start == 0 or start == end:
                    continue
                
                print(f"✅ Match for group '{group_name}': '{match.group()}'")
                if (keyword_string, rgb_color) not in data["found_words"]:
                    data["found_words"].append((keyword_string, rgb_color))

                # This is a simplified highlighting. For accurate highlighting,
                # we need to manipulate runs, which is more complex.
                # This approach will not work as intended with python-docx's run structure.
                # A better approach is to split the paragraph text and apply highlighting to runs.
    
    # Re-creating the paragraph with highlighted runs
    document.paragraphs[-1].clear() # Clear the plain text paragraph
    
    matches = []
    for group_name, data in keyword_groups.items():
        for keyword_string, pattern, rgb_color in data["patterns"]:
            for match in pattern.finditer(full_text):
                matches.append((match.start(), match.end(), rgb_color))

    matches.sort()

    last_end = 0
    p = document.paragraphs[-1]
    for start, end, color in matches:
        if start > last_end:
            p.add_run(full_text[last_end:start])
        
        run = p.add_run(full_text[start:end])
        font = run.font
        # Set text color to black as requested
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Set highlight color (background)
        font.highlight_color = get_closest_wd_color_index(color)

        last_end = end
    
    if last_end < len(full_text):
        p.add_run(full_text[last_end:])

    insert_summary_table(document, keyword_groups)

    document.save(OUTPUT_FILE)
    print(f'Document created: {os.path.abspath(OUTPUT_FILE)}')
    return keyword_groups

def insert_summary_table(document, summary_data):
    document.add_page_break()
    document.add_heading('Match Summary', level=1)
    
    rows = len(summary_data) + 1
    cols = 3
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # Headers
    headers = ["Sales Approach", "Found Words", "Status"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Data
    for r, (group, data) in enumerate(summary_data.items(), 1):
        table.cell(r, 0).text = group
        
        # Apply highlighting to found words in the summary table
        cell = table.cell(r, 1)
        p_cell = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p_cell.clear() # Clear existing text if any

        for i, (word, rgb_color) in enumerate(data["found_words"]):
            if i > 0:
                p_cell.add_run(", ") # Add comma and space separator
            run = p_cell.add_run(word)
            font = run.font
            font.color.rgb = RGBColor(0, 0, 0) # Text color black
            font.highlight_color = get_closest_wd_color_index(rgb_color)
        
        status = "Found" if data["found_words"] else "Not Found"
        table.cell(r, 2).text = status

def print_summary_table(summary_data):
    print("\n=== Match Summary by Group ===")
    if not summary_data:
        print("No keyword groups found.")
        return

    print(f"{'Sales Approach':<30} | {'Found Words':<40} | {'Status'}")
    print("-" * 80)
    for group, data in summary_data.items():
        status = "Found" if data["found_words"] else "Not Found"
        found_words_str = ", ".join([word for word, color in data["found_words"]])
        print(f"{group:<30} | {found_words_str:<40} | {status}")
    print("-" * 80)

# === Run ===
if __name__ == '__main__':
    full_text = load_transcript(TRANSCRIPT_FILE)
    keyword_groups = load_keyword_patterns(KEYWORDS_FILE)
    summary_result = create_docx_and_highlight(full_text, keyword_groups)
    print_summary_table(summary_result)
