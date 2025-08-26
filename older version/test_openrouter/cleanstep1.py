import csv  # นำเข้าโมดูล csv สำหรับอ่านไฟล์
import difflib
import json  # Import json module for JSONDecodeError
import logging
import os
import re  # เพิ่มเพื่อใช้งาน regex ลบ [Unrecognized]
import sys  # Import sys เพื่อใช้ในการกำหนด StreamHandler
import time  # Import time module for rate limiting

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from dotenv import load_dotenv
from openai import OpenAI

# --- Logging Configuration ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)  # แสดงผลใน Terminal
    ]
)

# --- Configuration ---
# Load environment variables from .env file
logging.info("Loading environment variables...")
load_dotenv()

# File and model settings
INPUT_FILE = r'transcript/debit/L_สาขาเดอะมอลล์งามวงศ์วาน.txt'
OUTPUT_FILE = r'transcript/debit_cleaned/cleaned_text_transcript_L_สาขาเดอะมอลล์งามวงศ์วาน.docx'
CLEANED_TEXT_OUTPUT_FILE = r'transcript/debit_cleaned/cleaned_text_transcript_L_สาขาเดอะมอลล์งามวงศ์วาน.txt'
MODEL_NAME = "google/gemma-3-27b-it:free"
KEYWORD_FILE = r'keywords/debit_card.csv'

# Get the API key from environment variables
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

# Define a chunk size for processing large texts
CHUNK_SIZE = 500  # characters

# --- Setup ---
# Ensure the output directory exists
output_dir = os.path.dirname(OUTPUT_FILE)
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Check for API key
if not OPENROUTER_API_KEY or OPENROUTER_API_KEY == "your_openrouter_api_key_here":
    logging.error("OpenRouter API key not found or is not set. Please update your .env file.")
    raise ValueError("OpenRouter API key not found or is not set. Please update your .env file.")
else:
    logging.info("OpenRouter API key loaded successfully.")

# Initialize the OpenAI client for OpenRouter
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=OPENROUTER_API_KEY,
)

# === Utility: ลบ [Unrecognized] ออกให้หมด โดยไม่แทนที่ด้วยคำอื่น ===
def remove_unrecognized(text: str) -> str:
    """
    ลบ token ที่เป็น [Unrecognized] ออกจากข้อความทั้งหมด
    รองรับรูปแบบที่มีช่องว่างในวงเล็บ และไม่สนตัวพิมพ์เล็ก/ใหญ่
    """
    # ลบ [Unrecognized] ทั้งหมด
    text = re.sub(r'\[\s*Unrecognized\s*\]', '', text, flags=re.IGNORECASE)
    # จัดช่องว่างให้เรียบร้อยหลังการลบ
    text = re.sub(r'\s{2,}', ' ', text).strip()
    return text


def read_keywords_from_csv(file_path: str) -> list[str]:
    """
    Reads a list of keywords from a CSV file, focusing on specific columns.
    """
    keywords = []
    try:
        with open(file_path, mode='r', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            # อ่านข้อมูลทั้งหมดจาก CSV
            rows = list(reader)
            if not rows:
                logging.warning("Keyword file is empty.")
                return []
            
            # สมมติว่าแถวแรกคือ Header และคอลัมน์ที่ต้องการคือ Word1 ถึง Word6
            headers = rows[0]
            # หา index ของคอลัมน์ Word1 ถึง Word6
            column_indices = []
            for i, header in enumerate(headers):
                if header.strip().lower() in [f"word{j}" for j in range(1, 7)]:
                    column_indices.append(i)

            if not column_indices:
                logging.error("Could not find Word1-Word6 columns in the CSV file.")
                return []

            # อ่านข้อมูลจากคอลัมน์ที่ต้องการเท่านั้น
            for row in rows[1:]:
                for index in column_indices:
                    if index < len(row):
                        word = row[index].strip()
                        if word:
                            keywords.append(word)

        # กรองคำที่ซ้ำกันออกไป และลบช่องว่างหัวท้าย
        keywords = [k for k in keywords if k]  # กรองคำว่างออก
        keywords = list(set(keywords))  # ลบคำซ้ำ
        
        logging.info(f"Successfully read {len(keywords)} keywords from {file_path}.")
        return keywords
    except FileNotFoundError:
        logging.error(f"Keyword file not found at {file_path}")
        return []
    except Exception as e:
        logging.error(f"An error occurred while reading the keyword file: {e}", exc_info=True)
        return []


def get_cleaned_text(text_chunk: str, keywords: list[str]) -> str:
    """
    Sends a text chunk to the specified OpenRouter model for cleaning,
    focusing only on words present in the keyword list.
    """
    logging.info(f"Sending text chunk to OpenRouter using model: {MODEL_NAME}...")
    logging.info(f"Text chunk length: {len(text_chunk)} characters.")
    logging.info(f"Text chunk snippet (first 500 chars): {text_chunk[:500]}")

    # Create the user prompt, embedding all instructions and keywords directly
    user_message_content = f"""Review the following Thai text for spelling mistakes and typos. Your task is to correct only words that are confusing or clearly misunderstood, and are misspellings of the provided keywords. Do not add new words, rephrase sentences, or make any changes beyond correcting these specific misspellings. The output should be the corrected text only, with no additional commentary. Also, remove all instances of the text "[Unrecognized]". The output must be plain text, with no markdown, special characters, or formatting whatsoever.

Keywords: {', '.join(keywords)}

Text to correct:
{text_chunk}
"""
    logging.info(f"Keywords added to user prompt: Keywords: {', '.join(keywords)[:100]}...")  # Log เฉพาะ snippet

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {
                    "role": "user",
                    "content": user_message_content
                }
            ],
            temperature=0.1,
            timeout=60.0,  # Add a timeout for the API call
        )
        logging.info("API call completed successfully.")
        cleaned_text = ""
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            cleaned_text = response.choices[0].message.content
            logging.info("Successfully received cleaned text from API.")
            logging.info(f"Received cleaned text snippet (first 200 chars): {cleaned_text[:200]}")
        else:
            logging.warning("API response did not contain expected message content.")
            logging.warning(f"Full response (if available): {response}")  # Log the raw response object

        logging.info("Exiting get_cleaned_text function for this chunk.")
        return cleaned_text
    except Exception as e:
        logging.error(f"An error occurred while calling the OpenRouter API for a chunk: {e}", exc_info=True)
        if isinstance(e, json.decoder.JSONDecodeError) and hasattr(e, 'doc'):
            logging.error(f"Raw API response text (causing JSONDecodeError): {e.doc}")
        return ""


def process_text_in_chunks(full_text: str, keywords: list[str]) -> str:
    """
    Splits the full text into chunks, sends each chunk to the API for cleaning,
    and concatenates the cleaned results.
    """
    logging.info(f"Processing full text in chunks. Total length: {len(full_text)} characters.")
    cleaned_chunks = []
    try:
        for i in range(0, len(full_text), CHUNK_SIZE):
            chunk = full_text[i:i + CHUNK_SIZE]
            logging.info(f"Processing chunk {int(i/CHUNK_SIZE) + 1} of {int(len(full_text)/CHUNK_SIZE) + 1}...")
            cleaned_chunk = get_cleaned_text(chunk, keywords)
            logging.info(f"Returned from get_cleaned_text for chunk {int(i/CHUNK_SIZE) + 1}.")
            cleaned_chunks.append(cleaned_chunk)
            time.sleep(5)  # Add a 5-second delay between API calls to avoid rate limiting
    except Exception as e:
        logging.error(f"An error occurred during chunk processing: {e}", exc_info=True)
        return ""
    
    final_cleaned_text = "".join(cleaned_chunks)
    logging.info(f"Finished processing all chunks. Total cleaned text length: {len(final_cleaned_text)} characters.")
    return final_cleaned_text


def create_highlighted_docx(original_text: str, cleaned_text: str, output_path: str, model_name: str):
    """
    Compares original and cleaned text, and creates a DOCX file
    highlighting only the corrected words (differences), not entire segments.
    """
    logging.info("Comparing texts and generating highlighted DOCX...")
    original_words = original_text.split()
    cleaned_words = cleaned_text.split()
    matcher = difflib.SequenceMatcher(None, original_words, cleaned_words)

    doc = Document()
    doc.add_heading(f'Cleaned Transcript with Highlights (Model: {model_name})', level=1)
    p = doc.add_paragraph()

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for word in cleaned_words[j1:j2]:
                p.add_run(word + ' ')
        elif tag == 'replace':
            orig = original_words[i1:i2]
            clean = cleaned_words[j1:j2]
            word_matcher = difflib.SequenceMatcher(None, orig, clean)
            for wtag, wi1, wi2, wj1, wj2 in word_matcher.get_opcodes():
                if wtag == 'equal':
                    for word in clean[wj1:wj2]:
                        p.add_run(word + ' ')
                else:
                    for word in clean[wj1:wj2]:
                        run = p.add_run(word + ' ')
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif tag == 'insert':
            for word in cleaned_words[j1:j2]:
                run = p.add_run(word + ' ')
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif tag == 'delete':
            continue  # ไม่ใส่คำที่ถูกลบจาก cleaned_text ลงในไฟล์

    doc.save(output_path)
    logging.info(f"Successfully created highlighted document: {output_path}")


# --- Main Execution ---
if __name__ == '__main__':
    logging.info(f"--- Starting process for {INPUT_FILE} ---")
    
    try:
        # 1. Read the keyword file
        keywords_list = read_keywords_from_csv(KEYWORD_FILE)
        
        # 2. Read the original file
        logging.info("Reading the original transcript...")
        with open(INPUT_FILE, 'r', encoding='utf-8') as f:
            original_text = f.read()
        logging.info("Successfully read the original transcript.")

        # 2.1 🔹 ลบ [Unrecognized] ออกจากต้นฉบับก่อนแบ่งชังก์ (กันเคสโดนตัดครึ่งตอนแบ่งชังก์)
        preprocessed_text = remove_unrecognized(original_text)

        # 3. Get the cleaned text from the API by processing in chunks (ใช้ข้อความที่ลบแล้ว)
        cleaned_text = process_text_in_chunks(preprocessed_text, keywords_list)
        
        # 3.1 🔹 กันพลาด: ลบ [Unrecognized] ออกจากผลลัพธ์อีกรอบ
        cleaned_text = remove_unrecognized(cleaned_text)
        
        # 4. Save the plain cleaned text to a separate file
        try:
            with open(CLEANED_TEXT_OUTPUT_FILE, 'w', encoding='utf-8') as f:
                f.write(cleaned_text)
            logging.info(f"Successfully saved plain cleaned text to: {CLEANED_TEXT_OUTPUT_FILE}")
        except Exception as e:
            logging.error(f"Could not write plain text file. Error: {e}")

        # 5. Create the highlighted DOCX, passing the model name
        #    ✅ ใช้ preprocessed_text (ที่ลบ [Unrecognized] แล้ว) ในการเทียบกับ cleaned_text
        create_highlighted_docx(preprocessed_text, cleaned_text, OUTPUT_FILE, MODEL_NAME)
        
        logging.info(f"--- Process complete. ---")
        
    except FileNotFoundError:
        logging.error(f"Input file not found at {INPUT_FILE}")
    except Exception as e:
        logging.critical(f"An unexpected error stopped the script: {e}", exc_info=True)
