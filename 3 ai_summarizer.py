import argparse  # สำหรับจัดการ argument ที่ส่งเข้ามาทาง command line
import glob  # สำหรับค้นหาไฟล์ตาม pattern
import json  # สำหรับทำงานกับข้อมูล JSON
import locale  # สำหรับตั้งค่าภาษาและรูปแบบท้องถิ่น
import logging  # สำหรับบันทึกข้อมูล log
import os  # สำหรับทำงานกับระบบปฏิบัติการ เช่น การจัดการไฟล์และ directory
import re  # สำหรับทำงานกับ Regular Expressions
import sys  # สำหรับเข้าถึงพารามิเตอร์และฟังก์ชันเฉพาะของระบบ
import time  # สำหรับฟังก์ชันที่เกี่ยวข้องกับเวลา
from datetime import datetime  # สำหรับทำงานกับวันที่และเวลา
from typing import (Any, Dict, List,  # สำหรับระบุประเภทข้อมูล (Type Hinting)
                    Optional, Tuple)

import docx  # สำหรับสร้างและแก้ไขไฟล์ .docx
from docx import Document  # คลาส Document จาก docx สำหรับสร้างเอกสาร
from docx.enum.section import (  # สำหรับกำหนดทิศทางและประเภทของ section ในเอกสาร
    WD_ORIENT, WD_SECTION)
from docx.enum.table import WD_ALIGN_VERTICAL  # สำหรับกำหนดการจัดแนวตาราง
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import \
    WD_ALIGN_PARAGRAPH  # สำหรับกำหนดการจัดแนวข้อความใน paragraph
from docx.oxml import OxmlElement  # สำหรับทำงานกับ XML ภายในไฟล์ .docx
from docx.oxml.ns import qn  # สำหรับจัดการ namespace ของ XML
from docx.shared import Inches, Mm, Pt, RGBColor  # สำหรับกำหนดหน่วยวัดและสี
from dotenv import load_dotenv  # สำหรับโหลดตัวแปรสภาพแวดล้อมจากไฟล์ .env
from openai import OpenAI  # สำหรับเชื่อมต่อกับ OpenAI API (หรือ OpenRouter)
from tqdm import tqdm  # สำหรับแสดง progress bar
from tqdm.contrib.logging import \
    _TqdmLoggingHandler as \
    TqdmLoggingHandler  # Handler สำหรับ logging ที่ทำงานร่วมกับ tqdm
from tqdm.contrib.logging import \
    logging_redirect_tqdm  # ฟังก์ชันสำหรับ redirect logging ให้แสดงผลเหนือ progress bar

# Set encoding
try:
    locale.setlocale(locale.LC_ALL, 'th_TH.utf-8') # ตั้งค่า locale เป็นภาษาไทย (UTF-8)
    logging.info("Successfully set locale to th_TH.utf-8") # บันทึก log เมื่อตั้งค่าสำเร็จ
except locale.Error as e:
    logging.warning(f"Could not set locale to th_TH.utf-8: {e}") # บันทึก log เมื่อตั้งค่า locale ไม่สำเร็จ
except Exception as e:
    logging.warning(f"Unexpected error setting locale: {e}") # บันทึก log เมื่อเกิดข้อผิดพลาดอื่น ๆ ในการตั้งค่า locale


# ทำให้ log แสดงสวยร่วมกับ tqdm (ขึ้นเหนือ progress bar)
logging.basicConfig(
    level=logging.INFO, # กำหนดระดับของ log ที่จะแสดง
    format="%(asctime)s - %(levelname)s - %(message)s", # รูปแบบการแสดงผล log
    handlers=[TqdmLoggingHandler()],  # << สำคัญ # ใช้ TqdmLoggingHandler เพื่อให้ log แสดงผลเหนือ progress bar
    force=True # บังคับให้ใช้การตั้งค่านี้
)

# กดเสียง lib ภายนอก (ไม่ให้ HTTP 200 แทรกกลาง bar)
logging.getLogger("httpx").setLevel(logging.WARNING) # ตั้งค่าระดับ log ของ httpx เป็น WARNING เพื่อลดข้อความที่ไม่จำเป็น
logging.getLogger("openai").setLevel(logging.WARNING) # ตั้งค่าระดับ log ของ openai เป็น WARNING เพื่อลดข้อความที่ไม่จำเป็น


load_dotenv() # Ensure environment variables are loaded at the start # โหลดตัวแปรสภาพแวดล้อมจากไฟล์ .env

# Status constants for scoring
PASS_STATUSES = {"พบ"} # สถานะที่ถือว่า "ผ่าน" ในการให้คะแนน

DECISION_GUIDE_TEXT = [ # ข้อความแนะนำสำหรับการตัดสินใจระดับคะแนนโดยรวม
    "ดีมาก: คะแนนรวมสูง (เช่น ≥90%) อ้างอิงจากขั้นตอนที่พนักงานต้องปฏิบัติตามกระบวนการขาย",
    "ปานกลาง: คะแนนรวมอยู่ในระดับพอใช้ (เช่น 60–89%)  อ้างอิงจากขั้นตอนที่พนักงานต้องปฏิบัติตามกระบวนการขาย",
    "ควรปรับปรุง: คะแนนรวมต่ำ (<60%)  อ้างอิงจากขั้นตอนที่พนักงานต้องปฏิบัติตามกระบวนการขาย",
]

EXCLUDED_STEPS_BY_PRODUCT = { # ขั้นตอนที่ไม่นำมาคิดคะแนนสำหรับแต่ละผลิตภัณฑ์
    "debit_card": {4, 5, 7}  # ขั้นตอนที่ไม่คิดคะแนนสำหรับ debit_card
}

# --- DOCX File Handling ---

def _repeat_header_on_each_page(row):
    """ตั้งค่าให้แถวตารางเป็น header ซ้ำทุกหน้า"""
    tr = row._tr # เข้าถึง element XML ของแถว
    trPr = tr.get_or_add_trPr() # เข้าถึงหรือเพิ่ม properties ของแถว
    tblHeader = OxmlElement('w:tblHeader') # สร้าง element สำหรับ header ตาราง
    trPr.append(tblHeader) # เพิ่ม header element เข้าไปใน properties ของแถว


def read_docx(file_path: str) -> str:
    """
    อ่านเนื้อหาทั้งหมดจากไฟล์ .docx รวมถึง paragraph และตาราง
    file_path: Path ของไฟล์ .docx
    return: ข้อความทั้งหมดที่อ่านได้จากไฟล์
    """
    try:
        doc = docx.Document(file_path) # เปิดเอกสาร .docx
        full_text = [] # list สำหรับเก็บข้อความทั้งหมด
        for para in doc.paragraphs: # วนอ่านทุก paragraph
            full_text.append(para.text) # เพิ่มข้อความจาก paragraph
        for table in doc.tables: # วนอ่านทุกตาราง
            for row in table.rows: # วนอ่านทุกแถวในตาราง
                for cell in row.cells: # วนอ่านทุก cell ในแถว
                    full_text.append(cell.text) # เพิ่มข้อความจาก cell
        return "\n".join(filter(None, full_text)) # รวมข้อความทั้งหมดเป็น string เดียว โดยกรองค่าว่างออก
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}") # บันทึก log เมื่อเกิดข้อผิดพลาดในการอ่านไฟล์
        return "" # คืนค่า string ว่างเปล่า

def truncate_transcript_if_needed(transcript: str, max_chars: int = 15000) -> str:
    """
    ตัดทอน transcript ถ้ายาวเกินขีดจำกัดที่กำหนด
    transcript: ข้อความ transcript ต้นฉบับ
    max_chars: จำนวนอักขระสูงสุดที่อนุญาต
    return: transcript ที่ถูกตัดทอนแล้ว (ถ้าจำเป็น)
    """
    if len(transcript) <= max_chars: # ถ้าความยาวไม่เกินขีดจำกัด
        return transcript # คืน transcript เดิม

    # ตัดจากตอนท้าย แต่พยายามตัดที่จุดจบประโยค
    truncated = transcript[:max_chars] # ตัด transcript ตามจำนวนอักขระสูงสุด
    last_period = truncated.rfind('.') # หาตำแหน่งของจุดสุดท้าย
    last_newline = truncated.rfind('\n') # หาตำแหน่งของขึ้นบรรทัดใหม่สุดท้าย

    cut_point = max(last_period, last_newline) # เลือกจุดตัดที่ใกล้เคียงจุดจบประโยคหรือขึ้นบรรทัดใหม่
    if cut_point > max_chars * 0.8:  # ถ้าจุดตัดไม่เสียข้อมูลมากเกินไป (อยู่ภายใน 80% ของ max_chars)
        truncated = truncated[:cut_point + 1] # ตัด transcript ที่จุดที่เหมาะสม

    return truncated + "\n\n[หมายเหตุ: เนื้อหาถูกตัดทอนเนื่องจากความยาวเกินขีดจำกัด]" # เพิ่มหมายเหตุเมื่อมีการตัดทอน

# --- Criteria Loading ---

def load_criteria_from_json(product_name: str, base_dir: str) -> Dict[str, Any]:
    """
    โหลดเกณฑ์การประเมินเฉพาะผลิตภัณฑ์จากไฟล์ JSON
    product_name: ชื่อผลิตภัณฑ์ (เช่น 'debit_card')
    base_dir: Directory ที่เก็บไฟล์ JSON เกณฑ์
    return: ข้อมูลเกณฑ์ที่โหลดจากไฟล์ JSON
    """
    script_dir = os.path.dirname(__file__) # Directory ของ script ปัจจุบัน
    path = os.path.join(script_dir, base_dir, f"{product_name}.json") # สร้าง path เต็มของไฟล์เกณฑ์
    if not os.path.exists(path): # ตรวจสอบว่าไฟล์มีอยู่จริงหรือไม่
        raise FileNotFoundError(f"Criteria file not found: {path}") # แจ้ง error ถ้าไม่พบไฟล์
    with open(path, "r", encoding="utf-8") as f: # เปิดไฟล์ JSON
        return json.load(f) # โหลดข้อมูล JSON

# --- AI Interaction & Prompting ---

def get_openai_client() -> OpenAI:
    """
    เริ่มต้นและคืนค่า OpenAI client สำหรับ OpenRouter
    return: Instance ของ OpenAI client
    """
    api_key = os.getenv("OPENROUTER_API_KEY") # ดึง API key จากตัวแปรสภาพแวดล้อม
    if not api_key: # ถ้าไม่มี API key
        raise ValueError("OPENROUTER_API_KEY environment variable not set.") # แจ้ง error
    return OpenAI(
        base_url="https://openrouter.ai/api/v1", # กำหนด base URL สำหรับ OpenRouter
        api_key=api_key, # กำหนด API key
        timeout=30.0, # กำหนด timeout สำหรับ request
    )

def is_rate_limit_error(e: Exception) -> bool:
    """
    ตรวจสอบว่า exception ที่เกิดขึ้นเป็น error จาก rate limit (HTTP 429) หรือไม่
    e: Exception ที่เกิดขึ้น
    return: True ถ้าเป็น rate limit error, False ถ้าไม่ใช่
    """
    return "429" in str(e) # ตรวจสอบว่ามี "429" ในข้อความ error หรือไม่

def extract_json_from_response(text: str) -> Optional[Dict]:
    """
    พยายามดึง JSON object ออกจากข้อความตอบกลับของ AI
    text: ข้อความตอบกลับจาก AI
    return: JSON object ที่ดึงได้ หรือ None ถ้าไม่พบ
    """
    if not text: # ถ้าข้อความเป็นค่าว่าง
        return None # คืนค่า None
    s = text.strip() # ลบช่องว่างที่หัวท้าย
    # ตัด code fences ถ้ามี
    s = re.sub(r"^```[a-zA-Z_]*\s*", "", s).strip() # ลบ code fence เปิด
    s = re.sub(r"\s*```$", "", s).strip() # ลบ code fence ปิด

    # 1) ลอง parse ตรง ๆ ก่อน
    try:
        obj = json.loads(s) # พยายาม parse เป็น JSON โดยตรง
        return obj # คืน JSON object ถ้าสำเร็จ
    except Exception:
        pass # ถ้าไม่สำเร็จ ให้ลองวิธีอื่น

    # 2) แตกทุก {...} เป็น candidates แล้วเลือกก้อนที่ "เข้ารูป" ที่สุด
    candidates = [m.group(0) for m in re.finditer(r"\{.*?\}", s, flags=re.DOTALL)] # หา JSON-like string ทั้งหมด
    best = None # เก็บ JSON object ที่ดีที่สุด
    best_score = -1 # คะแนนของ JSON object ที่ดีที่สุด

    for cand in candidates: # วนลูปใน candidate ทั้งหมด
        try:
            obj = json.loads(cand) # พยายาม parse candidate เป็น JSON
        except Exception:
            continue # ถ้าไม่สำเร็จ ข้ามไป candidate ถัดไป
        score = 0 # คะแนนสำหรับ candidate ปัจจุบัน
        steps = obj.get("steps") # ดึงค่า "steps"
        if isinstance(steps, list): # ถ้า "steps" เป็น list
            # ให้แต้มตามจำนวน items จริง
            score += len(steps) # เพิ่มคะแนนตามจำนวน steps
            # แต้มเพิ่มถ้ามี items ภายใน
            score += sum(1 for st in steps if isinstance(st, dict) and isinstance(st.get("items"), list) and len(st["items"]) > 0) # เพิ่มคะแนนถ้ามี items ภายใน steps
        if isinstance(obj.get("summary"), dict): # ถ้า "summary" เป็น dict
            score += 1 # เพิ่มคะแนน
        if score > best_score: # ถ้าคะแนนปัจจุบันดีกว่าคะแนนที่ดีที่สุด
            best = obj # อัปเดต best JSON object
            best_score = score # อัปเดต best score

    return best # คืน JSON object ที่ดีที่สุด

def call_ai_model(client: OpenAI, messages: List[Dict[str, Any]], model: str, max_retries: int = 3) -> Optional[Dict[str, Any]]:
    """
    เรียกโมเดล AI แบบเรียบง่าย พร้อมระบบ backoff สำหรับ error 429
    - ไม่พิมพ์เนื้อหาคำตอบลง log
    - ไม่สร้างไฟล์ raw_ai_response.txt
    client: OpenAI client
    messages: list ของข้อความสำหรับส่งให้ AI
    model: ชื่อโมเดล AI ที่จะใช้
    max_retries: จำนวนครั้งสูงสุดที่จะลองใหม่เมื่อเกิด error
    return: JSON object ที่ได้จาก AI หรือ None ถ้าไม่สำเร็จ
    """
    retries = 0 # จำนวนครั้งที่ลองใหม่
    wait_time = 1.0 # เวลารอเริ่มต้น
    while retries <= max_retries: # วนลูปตามจำนวนครั้งที่ลองใหม่สูงสุด
        try:
            logging.info(f"Calling model {model} (attempt {retries + 1})...") # บันทึก log การเรียกโมเดล
            response = client.chat.completions.create( # เรียก API ของ AI
                model=model, # กำหนดโมเดล
                messages=messages, # กำหนดข้อความ
                temperature=0.1, # กำหนดอุณหภูมิ (ความสุ่ม) ของการตอบกลับ
            )
            content = (response.choices[0].message.content or "").strip() # ดึงเนื้อหาการตอบกลับ
            logging.info(f"Model {model} returned {len(content)} characters") # บันทึก log ความยาวของเนื้อหา

            result = extract_json_from_response(content) # พยายามดึง JSON จากเนื้อหา
            if result: # ถ้าดึง JSON ได้
                result["_model_used"] = model # เพิ่มชื่อโมเดลที่ใช้
                logging.info(f"Successfully extracted JSON from {model}") # บันทึก log เมื่อดึง JSON สำเร็จ
                return result # คืนผลลัพธ์
            
            logging.warning(f"Model {model} returned non-JSON or invalid JSON. Content preview: {content[:200]}...") # บันทึก log เมื่อ AI ตอบกลับไม่ใช่ JSON ที่ถูกต้อง
            return None # คืนค่า None

        except Exception as e: # ดักจับ exception
            logging.error(f"Model {model} error (attempt {retries + 1}): {str(e)}") # บันทึก log เมื่อเกิด error
            if is_rate_limit_error(e) and retries < max_retries: # ถ้าเป็น rate limit error และยังไม่เกินจำนวนครั้งที่ลองใหม่
                logging.warning(f"429 from {model}. Retry in {wait_time:.1f}s...") # บันทึก log การรอเพื่อลองใหม่
                time.sleep(wait_time); wait_time *= 2; retries += 1 # รอและเพิ่มเวลารอ
                continue # ลองใหม่
            return None # คืนค่า None เมื่อไม่สามารถประมวลผลได้

def retry_force_json(client: OpenAI, base_prompt: str, model_list: List[str]) -> Optional[Dict]:
    """
    พยายามเรียก AI อีกครั้งด้วยคำสั่งที่เข้มงวดขึ้น เพื่อให้คืนค่าเป็น JSON object เดียวเท่านั้น
    client: OpenAI client
    base_prompt: prompt พื้นฐาน
    model_list: list ของโมเดล AI ที่จะลองใช้
    return: JSON object ที่ได้จาก AI หรือ None ถ้าไม่สำเร็จ
    """
    strict = (
        "คำสั่งสำคัญ: ตอบเป็น JSON object เดียวเท่านั้น ห้ามมีข้อความอื่นหรือ code fence ใด ๆ นอกเหนือจาก JSON\n"
        "ห้ามใส่ตัวอย่าง ไม่ต้องอธิบาย\n"
        "ต้องมี keys: metadata, steps, summary ตาม schema ที่กำหนด"
    ) # คำสั่งเข้มงวดสำหรับ AI
    strict_prompt = strict + "\n\n" + base_prompt # รวมคำสั่งเข้มงวดกับ prompt พื้นฐาน
    for m in model_list: # วนลูปใน list ของโมเดล
        resp = call_ai_model(client, [{"role": "user", "content": strict_prompt}], m) # เรียก AI ด้วย prompt เข้มงวด
        if resp and isinstance(resp.get("steps"), list) and len(resp["steps"]) > 0: # ถ้าได้ผลลัพธ์และมี "steps" ที่ถูกต้อง
            resp["_model_used"] = m # เพิ่มชื่อโมเดลที่ใช้
            return resp # คืนผลลัพธ์
    return None # คืนค่า None ถ้าไม่สำเร็จ

def build_comprehensive_prompt(transcript: str, criteria_data: Dict[str, Any]) -> str:
    """
    สร้าง prompt ที่ครอบคลุมสำหรับ AI เพื่อให้ทำงานทั้งหมดในครั้งเดียว
    transcript: ข้อความ transcript ที่จะให้ AI วิเคราะห์
    criteria_data: ข้อมูลเกณฑ์การประเมิน
    return: prompt ที่สร้างขึ้น
    """
    criteria_ref = [] # list สำหรับเก็บข้อมูลอ้างอิงเกณฑ์
    for step_no_str, items in criteria_data["criteria_by_step"].items(): # วนลูปในเกณฑ์แต่ละขั้นตอน
        step_no = int(step_no_str) # แปลงเลขขั้นตอนเป็น int
        for idx, item in enumerate(items, 1): # วนลูปในเกณฑ์ย่อยแต่ละรายการ
            name = item.get("name", str(item)) # ดึงชื่อเกณฑ์ย่อย
            criteria_ref.append({"id": f"{step_no}.{idx}", "step": step_no, "name": name}) # เพิ่มข้อมูลเกณฑ์ย่อย

    prompt = f"""
คุณคือผู้ตรวจสอบคุณภาพ (QA) และผู้จัดการรายงานที่เชี่ยวชาญสูง
หน้าที่ของคุณคือทำ 2 อย่างในครั้งเดียว:
1.  วิเคราะห์บทสนทนา (Transcript) อย่างละเอียด และประเมินตามเกณฑ์ที่กำหนด
2.  เขียนสรุปผลการประเมินในเชิงคุณภาพ เป็นภาษาที่สละสลวย

**คุณต้องตอบกลับเป็น JSON object ที่สมบูรณ์เพียง object เดียวเท่านั้น ห้ามมีข้อความอื่นใดนอกเหนือจาก JSON object และห้ามเพิ่ม field อื่นๆ นอกเหนือจากที่กำหนดไว้ในโครงสร้าง JSON ด้านล่างนี้เด็ดขาด**

### 1. บทสนทนาที่ต้องวิเคราะห์ (Transcript)
```text
{transcript}
```

### 2. โครงสร้าง JSON ที่คุณต้องตอบกลับ (Output Format)
คุณต้องกรอกข้อมูลตามโครงสร้างนี้ให้ครบถ้วน โดยอ้างอิงจาก Transcript และเกณฑ์ที่ให้ไว้เท่านั้น:
```json
{{
  "metadata": {{
    "date": "<string or 'ไม่ระบุ'>",
    "branch": "<string or 'ไม่ระบุ'>"
  }},
  "steps": [
    {{
      "step": "<int>",
      "title": "<string>",
      "items": [
        {{
          "id": "<string, e.g., '1.1'>",
          "name": "<string, from CRITERIA_REFERENCE>",
          "status": "พบ|ไม่พบ|พบบางส่วน",
          "reason": "<string, provide a brief reason ONLY if status is 'ไม่พบ' or 'พบบางส่วน'>",
          "evidence": {{
            "exact": "<string, a short quote from the transcript>",
            "sentence": "<string, the full sentence containing the evidence>",
            "offset": ["<string, e.g., '[34..62]'>"]
          }}
        }}
      ]
    }}
  ],
  "summary": {{
    "compliance": {{
      "do": ["<list of observed positive behaviors based on Market Conduct 'Do' list>"],
      "dont": ["<list of observed violations based on Market Conduct 'Don't' list>"]
    }},
    "narrative_reason": "<string, เขียนคำอธิบายเชิงเหตุผล 1 ย่อหน้า ว่าทำไมการสนทนานี้จึงถูกประเมินในระดับคะแนนที่ได้ โดยสังเคราะห์จากภาพรวม ไม่ใช่แค่การลิสต์คะแนน แต่ให้เล่าเรื่องว่าพนักงานทำอะไรได้ดี และขาดตกบกพร่องในขั้นตอนสำคัญใดบ้าง>",
    "strengths": ["<list of 2-3 key strengths>"],
    "improvements": ["<list of 2-3 key areas for improvement>"]
  }}
}}
```

### 3. เกณฑ์การประเมินและข้อมูลอ้างอิง (CRITERIA_REFERENCE)
- **ขั้นตอนและหัวข้อ:** {json.dumps(criteria_data["steps"], ensure_ascii=False)}
- **เกณฑ์ย่อย (ห้ามเพิ่มหรือลบรายการ):** {json.dumps(criteria_ref, ensure_ascii=False, indent=2)}
- **Market Conduct (Do):** {json.dumps(criteria_data.get("compliance_rules", {}).get("Do", []), ensure_ascii=False)}
- **Market Conduct (Don't):** {json.dumps(criteria_data.get("compliance_rules", {}).get("Dont", []), ensure_ascii=False)}

**กฎสำคัญ:**
- **ยึดตาม Transcript เท่านั้น:** ห้ามคาดเดาหรือเสริมข้อมูลที่ไม่มีในบทสนทนา
- **สถานะ:** ทุกเกณฑ์ต้องมีสถานะอย่างใดอย่างหนึ่ง: "พบ", "ไม่พบ", "พบบางส่วน"
- **หลักฐาน (Evidence):** ต้องมาจาก Transcript จริงๆ
- **สรุปผล:** ต้องเขียนสรุป `narrative_reason`, `strengths`, และ `improvements` ให้ครบถ้วนและมีคุณภาพ
- **ห้ามเพิ่ม field อื่นๆ นอกเหนือจากที่กำหนดในโครงสร้าง JSON Output Format เด็ดขาด**
- **หากไม่มีข้อมูลสำหรับ field ใดๆ ให้ใช้ค่าว่าง (empty string for string, empty list for list) แทนการละเว้น field นั้นๆ**
"""
    return prompt # คืน prompt ที่สร้างขึ้น

def process_single_file_with_retries(transcript: str, criteria_data: Dict, client: OpenAI, models: List[str], max_attempts: int = 3):
    """
    ประมวลผลไฟล์เดียวด้วยการ retry หลายครั้งเมื่อเกิดปัญหา
    transcript: ข้อความ transcript ที่จะประมวลผล
    criteria_data: ข้อมูลเกณฑ์การประเมิน
    client: OpenAI client
    models: list ของโมเดล AI ที่จะลองใช้
    max_attempts: จำนวนครั้งสูงสุดที่จะลองประมวลผล
    return: รายงานผลลัพธ์จาก AI หรือ None ถ้าไม่สำเร็จ
    """
    for attempt in range(max_attempts): # วนลูปตามจำนวนครั้งที่ลองสูงสุด
        logging.info(f"Processing attempt {attempt + 1}/{max_attempts}") # บันทึก log การพยายามประมวลผล

        # ถ้าพยายามมากกว่า 1 ครั้ง ให้ตัด transcript ลง
        current_transcript = transcript # กำหนด transcript ปัจจุบัน
        if attempt > 0: # ถ้าเป็นการลองครั้งที่ 2 ขึ้นไป
            current_transcript = truncate_transcript_if_needed(transcript, 12000 - (attempt * 2000)) # ตัดทอน transcript ให้สั้นลง
            logging.info(f"Truncated transcript to {len(current_transcript)} characters") # บันทึก log ความยาว transcript ที่ถูกตัดทอน

        prompt = build_comprehensive_prompt(current_transcript, criteria_data) # สร้าง prompt

        for model in models: # วนลูปใน list ของโมเดล
            report = call_ai_model(client, [{"role": "user", "content": prompt}], model) # เรียก AI model
            if report: # ถ้าได้รายงานผลลัพธ์
                steps_count = sum(len(st.get("items", [])) for st in report.get("steps", []) if isinstance(st, dict)) # นับจำนวน items ที่ถูกวิเคราะห์
                if steps_count > 0:  # ตรวจสอบว่ามีการวิเคราะห์จริง
                    logging.info(f"Successfully processed with {steps_count} analyzed items") # บันทึก log เมื่อประมวลผลสำเร็จ
                    return report # คืนรายงานผลลัพธ์

        # หน่วงเวลาก่อน retry
        if attempt < max_attempts - 1: # ถ้ายังไม่ถึงจำนวนครั้งที่ลองสูงสุด
            wait_time = 10 * (attempt + 1) # คำนวณเวลารอ
            logging.info(f"Waiting {wait_time} seconds before retry...") # บันทึก log การรอ
            time.sleep(wait_time) # รอตามเวลาที่กำหนด

    logging.error(f"Failed to process after {max_attempts} attempts") # บันทึก log เมื่อประมวลผลไม่สำเร็จหลังจากลองหลายครั้ง
    return None # คืนค่า None

# --- Scoring and Analysis ---

def compute_scores(report: Dict[str, Any], product_name: str = None) -> Tuple[Dict[str, Any], Dict[int, Tuple[int, int]]]:
    """
    คำนวณคะแนนแต่ละขั้นตอนและคะแนนรวม แล้วเพิ่มลงในรายงาน
    report: รายงานผลลัพธ์จาก AI
    product_name: ชื่อผลิตภัณฑ์ (ใช้สำหรับยกเว้นขั้นตอนที่ไม่คิดคะแนน)
    return: รายงานที่เพิ่มคะแนนแล้ว และคะแนนรายขั้นตอน
    """
    per_step_scores: Dict[int, Tuple[int, int]] = {} # dictionary สำหรับเก็บคะแนนรายขั้นตอน (ผ่าน, ทั้งหมด)
    total_pass = 0 # คะแนนรวมที่ผ่าน
    total_all = 0 # คะแนนรวมทั้งหมด

    # ดึงขั้นตอนที่ไม่คิดคะแนนสำหรับ product นี้
    excluded_steps = EXCLUDED_STEPS_BY_PRODUCT.get(product_name, set()) # ดึงชุดของขั้นตอนที่ถูกยกเว้น

    for step in report.get("steps", []): # วนลูปในแต่ละขั้นตอนในรายงาน
        step_no = int(step["step"]) # เลขที่ขั้นตอน
        step_pass = 0 # คะแนนที่ผ่านสำหรับขั้นตอนนี้
        step_all = 0 # คะแนนทั้งหมดสำหรับขั้นตอนนี้

        # คำนวณคะแนนของแต่ละขั้นตอน
        for item in step.get("items", []): # วนลูปในแต่ละ item ย่อยในขั้นตอน
            status = (item.get("status") or "").strip() # สถานะของ item
            step_all += 1 # เพิ่มจำนวน item ทั้งหมด
            if status in PASS_STATUSES: # ถ้าสถานะเป็น "ผ่าน"
                step_pass += 1 # เพิ่มคะแนนที่ผ่าน

        per_step_scores[step_no] = (step_pass, step_all) # เก็บคะแนนรายขั้นตอน

        # ถ้าเป็นขั้นตอนที่ไม่คิดคะแนน ให้แสดงเป็น 0/0 และเพิ่มหมายเหตุ
        if step_no in excluded_steps: # ถ้าขั้นตอนนี้ถูกยกเว้น
            step["score"] = f"0/0 (ไม่นำข้อนี้มาคิดคะแนน)" # กำหนด score เป็น "0/0 (ไม่นำข้อนี้มาคิดคะแนน)"
        else:
            step["score"] = f"{step_pass}/{step_all}" # กำหนด score เป็น "ผ่าน/ทั้งหมด"
            # เฉพาะขั้นตอนที่คิดคะแนนเท่านั้นที่จะนำมารวม
            total_pass += step_pass # เพิ่มคะแนนที่ผ่านรวม
            total_all += step_all # เพิ่มคะแนนทั้งหมดรวม

    percent = round(100.0 * total_pass / total_all, 1) if total_all > 0 else 0.0 # คำนวณเปอร์เซ็นต์รวม

    if "summary" not in report or not isinstance(report["summary"], dict): # ถ้าไม่มี summary หรือ summary ไม่ใช่ dict
        report["summary"] = {} # สร้าง summary เปล่า

    report["summary"]["_score_total"] = f"{total_pass}/{total_all}" # เพิ่มคะแนนรวมทั้งหมดใน summary
    report["summary"]["_score_percent"] = f"{percent}%" # เพิ่มเปอร์เซ็นต์รวมใน summary

    return report, per_step_scores # คืนรายงานและคะแนนรายขั้นตอน

def normalize_report(report: Dict[str, Any], criteria_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    ปรับโครงสร้างรายงานให้เป็นมาตรฐาน โดยเติมข้อมูลที่ขาดหายไปจากเกณฑ์
    report: รายงานผลลัพธ์จาก AI
    criteria_data: ข้อมูลเกณฑ์การประเมิน
    return: รายงานที่ถูกปรับโครงสร้างแล้ว
    """
    if not isinstance(report, dict): # ถ้า report ไม่ใช่ dict
        report = {} # สร้าง report เปล่า

    # เติม metadata/summary เปล่า ๆ
    report.setdefault("metadata", {"date": "ไม่ระบุ", "branch": "ไม่ระบุ"}) # กำหนด metadata เริ่มต้นถ้าไม่มี
    report.setdefault("summary", { # กำหนด summary เริ่มต้นถ้าไม่มี
        "compliance": {"do": [], "dont": []},
        "narrative_reason": "",
        "strengths": [],
        "improvements": []
    })

    steps = report.get("steps") # ดึง steps จากรายงาน
    if not (isinstance(steps, list) and len(steps) > 0): # ถ้า steps ไม่ใช่ list หรือเป็น list ว่าง
        # สร้าง steps เปล่าจาก criteria
        steps_new = [] # list สำหรับ steps ใหม่
        for step_no_str, items in criteria_data.get("criteria_by_step", {}).items(): # วนลูปในเกณฑ์แต่ละขั้นตอน
            step_no = int(step_no_str) # เลขที่ขั้นตอน
            # หา title
            title = next((s.split(')')[1].strip() for s in criteria_data.get("steps", [])
                    if int(s.split(')')[0].strip()) == step_no), f"ขั้นตอน {step_no}") # หา title ของขั้นตอน
            new_items = [] # list สำหรับ items ย่อยใหม่
            for idx, item in enumerate(items, 1): # วนลูปในเกณฑ์ย่อยแต่ละรายการ
                name = item.get("name", str(item)) # ดึงชื่อเกณฑ์ย่อย
                new_items.append({ # เพิ่ม item ย่อยใหม่
                    "id": f"{step_no}.{idx}",
                    "name": name,
                    "status": "",
                    "reason": "",
                    "evidence": {"exact": "", "sentence": "", "offset": []}
                })
            steps_new.append({"step": step_no, "title": title, "items": new_items}) # เพิ่มขั้นตอนใหม่
        steps_new.sort(key=lambda x: int(x["step"])) # เรียงลำดับ steps ตามเลขที่ขั้นตอน
        report["steps"] = steps_new # กำหนด steps ใหม่ให้กับรายงาน

    return report # คืนรายงานที่ถูกปรับโครงสร้างแล้ว

def classify_overall(report: Dict[str, Any]) -> str:
    """
    จัดระดับการประเมินโดยรวมจากเปอร์เซ็นต์คะแนน
    report: รายงานผลลัพธ์จาก AI
    return: ระดับการประเมิน (ดีมาก, ปานกลาง, ควรปรับปรุง)
    """
    summary = report.get("summary", {}) # ดึง summary จากรายงาน
    percent_str = str(summary.get("_score_percent", "0%")).replace('%', '').strip() # ดึงเปอร์เซ็นต์คะแนนเป็น string และลบ '%'
    try:
        percent = float(percent_str) # แปลงเปอร์เซ็นต์เป็น float
    except Exception:
        percent = 0.0 # ถ้าแปลงไม่ได้ กำหนดเป็น 0.0

    # ดีมาก >= 90, ปานกลาง 60–89, ควรปรับปรุง < 60
    if percent >= 90: # ถ้าเปอร์เซ็นต์มากกว่าหรือเท่ากับ 90
        return "ดีมาก" # คืน "ดีมาก"
    elif 60 <= percent < 90: # ถ้าเปอร์เซ็นต์อยู่ระหว่าง 60 ถึง 89
        return "ปานกลาง" # คืน "ปานกลาง"
    else: # ถ้าเปอร์เซ็นต์น้อยกว่า 60
        return "ควรปรับปรุง" # คืน "ควรปรับปรุง"

# --- DOCX Report Rendering ---

def _apply_run_font(run, *, name="Tahoma", size_pt=11, bold=False):
    """
    ตั้งฟอนต์/ขนาด/หนา ให้ run รองรับภาษาไทยด้วย eastAsia
    run: run object ใน docx
    name: ชื่อฟอนต์
    size_pt: ขนาดฟอนต์ (point)
    bold: ตัวหนาหรือไม่
    """
    run.bold = bold # กำหนดตัวหนา
    run.font.size = Pt(size_pt) # กำหนดขนาดฟอนต์
    run.font.name = name # กำหนดชื่อฟอนต์
    # สำคัญสำหรับภาษาไทย
    if run._element.rPr is not None: # ถ้ามี properties ของ run
        rFonts = run._element.rPr.rFonts # เข้าถึง rFonts
        if rFonts is None: # ถ้าไม่มี rFonts
            from docx.oxml import OxmlElement  # import OxmlElement
            rFonts = OxmlElement('w:rFonts') # สร้าง rFonts element
            run._element.rPr.append(rFonts) # เพิ่ม rFonts เข้าไป
        rFonts.set(qn('w:eastAsia'), name) # ตั้งค่าฟอนต์สำหรับ East Asia (ภาษาไทย)
        rFonts.set(qn('w:ascii'), name) # ตั้งค่าฟอนต์สำหรับ ASCII
        rFonts.set(qn('w:hAnsi'), name) # ตั้งค่าฟอนต์สำหรับ High ANSI

def _paragraph_with_text(doc, text, *, size_pt=11, bold=False, before=0, after=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    สร้าง paragraph พร้อมตั้งฟอนต์ตามต้องการ
    doc: Document object
    text: ข้อความใน paragraph
    size_pt: ขนาดฟอนต์
    bold: ตัวหนาหรือไม่
    before: ระยะห่างก่อน paragraph (point)
    after: ระยะห่างหลัง paragraph (point)
    align: การจัดแนวข้อความ
    return: paragraph object
    """
    p = doc.add_paragraph() # เพิ่ม paragraph ใหม่
    p.paragraph_format.space_before = Pt(before) # กำหนดระยะห่างก่อน paragraph
    p.paragraph_format.space_after = Pt(after) # กำหนดระยะห่างหลัง paragraph
    p.alignment = align # กำหนดการจัดแนวข้อความ
    r = p.add_run(text or "") # เพิ่ม run และข้อความ
    _apply_run_font(r, size_pt=size_pt, bold=bold) # ตั้งค่าฟอนต์
    return p # คืน paragraph object

def _set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    """
    Helper สำหรับจัดรูปแบบข้อความภายใน cell ของตาราง (Tahoma 11 โดยค่าเริ่มต้น)
    cell: cell object ในตาราง
    text: ข้อความที่จะใส่ใน cell
    bold: ตัวหนาหรือไม่
    align: การจัดแนวข้อความ
    size: ขนาดฟอนต์
    """
    cell.text = "" # ล้างข้อความเดิมใน cell
    p = cell.paragraphs[0] # เข้าถึง paragraph แรกใน cell
    p.alignment = align # กำหนดการจัดแนวข้อความ
    run = p.add_run(str(text or "")) # เพิ่ม run และข้อความ
    _apply_run_font(run, size_pt=size, bold=bold) # ตั้งค่าฟอนต์
    p.paragraph_format.space_before = Pt(1) # กำหนดระยะห่างก่อน paragraph
    p.paragraph_format.space_after = Pt(1) # กำหนดระยะห่างหลัง paragraph

def _set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    """
    Helper สำหรับจัดรูปแบบข้อความภายใน cell ของตาราง (Tahoma 11 โดยค่าเริ่มต้น)
    cell: cell object ในตาราง
    text: ข้อความที่จะใส่ใน cell
    bold: ตัวหนาหรือไม่
    align: การจัดแนวข้อความ
    size: ขนาดฟอนต์
    """
    cell.text = "" # ล้างข้อความเดิมใน cell
    p = cell.paragraphs[0] # เข้าถึง paragraph แรกใน cell
    p.alignment = align # กำหนดการจัดแนวข้อความ
    run = p.add_run(str(text or "")) # เพิ่ม run และข้อความ
    _apply_run_font(run, size_pt=size, bold=bold) # ตั้งค่าฟอนต์
    p.paragraph_format.space_before = Pt(1) # กำหนดระยะห่างก่อน paragraph
    p.paragraph_format.space_after = Pt(1) # กำหนดระยะห่างหลัง paragraph

def _shrink_font(cell, pt):
    """
    ลดขนาดฟอนต์ของข้อความใน cell
    cell: cell object ในตาราง
    pt: ขนาดฟอนต์ใหม่ (point)
    """
    for paragraph in cell.paragraphs: # วนลูปในทุก paragraph ใน cell
        for run in paragraph.runs: # วนลูปในทุก run ใน paragraph
            _apply_run_font(run, size_pt=pt) # ตั้งค่าขนาดฟอนต์ใหม่

# ---------- helpers for layout & font ----------

def _ensure_tahoma_styles(doc, body_size=11, h1_size=16):
    """
    เซ็ตฟอนต์เริ่มต้นทั้งเอกสารให้เป็น Tahoma (เผื่อภาษาไทย)
    doc: Document object
    body_size: ขนาดฟอนต์สำหรับเนื้อหาปกติ
    h1_size: ขนาดฟอนต์สำหรับ Heading 1
    """
    st = doc.styles['Normal'] # เข้าถึง style 'Normal'
    st.font.name = "Tahoma" # กำหนดชื่อฟอนต์
    st.font.size = Pt(body_size) # กำหนดขนาดฟอนต์
    st._element.rPr.rFonts.set(qn('w:eastAsia'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ East Asia
    st._element.rPr.rFonts.set(qn('w:ascii'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ ASCII
    st._element.rPr.rFonts.set(qn('w:hAnsi'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ High ANSI

    h1 = doc.styles['Heading 1'] # เข้าถึง style 'Heading 1'
    h1.font.name = "Tahoma" # กำหนดชื่อฟอนต์
    h1.font.size = Pt(h1_size) # กำหนดขนาดฟอนต์
    h1.font.bold = True # กำหนดตัวหนา
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ East Asia
    h1._element.rPr.rFonts.set(qn('w:ascii'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ ASCII
    h1._element.rPr.rFonts.set(qn('w:hAnsi'), "Tahoma") # ตั้งค่าฟอนต์สำหรับ High ANSI

def _printable_width(section):
    """
    คำนวณความกว้างพื้นที่พิมพ์ของ section (หน่วย: นิ้ว)
    section: section object
    return: ความกว้างพื้นที่พิมพ์
    """
    # ความกว้างพื้นที่พิมพ์ (นิ้ว) = ความกว้างหน้า - margin ซ้าย/ขวา
    return (section.page_width - section.left_margin - section.right_margin) / Inches(1.0) # คำนวณและคืนค่า

def _apply_widths_fit(tbl, widths_in, section, safety=0.88):
    """
    บังคับความกว้างรวมของตารางให้ <= พื้นที่พิมพ์ในหน้า (page_width - margins) * safety
    tbl: table object
    widths_in: รายการสัดส่วนความกว้างที่ต้องการ (หน่วย 'นิ้ว' ก่อนสเกล) เช่น [1.2, 0.6, 1.3, 1.5, 3.2, 0.6]
    section: section object
    safety: ค่าเผื่อระยะเพื่อไม่ให้ล้น (0.86–0.92)
    """
    page_in = section.page_width.inches # ความกว้างหน้ากระดาษ (นิ้ว)
    left_in = section.left_margin.inches # ระยะขอบซ้าย (นิ้ว)
    right_in = section.right_margin.inches # ระยะขอบขวา (นิ้ว)

    printable_in = max(0.1, page_in - left_in - right_in) * safety # คำนวณความกว้างที่พิมพ์ได้จริง
    sum_desired = sum(widths_in) # ผลรวมของความกว้างที่ต้องการ
    scale = printable_in / sum_desired if sum_desired > 0 else 1.0 # คำนวณ scale factor

    # ตั้งความกว้างรวมตาราง
    tbl.autofit = False # ปิด autofit
    tbl.allow_autofit = False if hasattr(tbl, "allow_autofit") else None # ปิด allow_autofit
    tbl.preferred_width = Inches(printable_in) # กำหนดความกว้างที่ต้องการ

    # เซ็ตความกว้างแต่ละคอลัมน์
    for j, w in enumerate(widths_in): # วนลูปในความกว้างแต่ละคอลัมน์
        col_w = max(0.2, w * scale)  # กันคอลัมน์เล็กเกิน # คำนวณความกว้างคอลัมน์จริง
        tbl.columns[j].width = Inches(col_w) # กำหนดความกว้างคอลัมน์
        for cell in tbl.columns[j].cells: # วนลูปใน cell ของคอลัมน์
            cell.width = Inches(col_w) # กำหนดความกว้าง cell

def _set_table_cell_margins(table, left=50, right=50, top=30, bottom=30):
    """
    ปรับระยะขอบภายในเซลล์ (หน่วย: twips) ลด padding เพื่อให้ตัวอักษรมีพื้นที่เพิ่ม
    ค่า default ค่อนข้างเล็กอยู่แล้ว แต่อาจช่วยได้มากเมื่อข้อความยาว
    table: table object
    left, right, top, bottom: ระยะขอบ (twips)
    """
    tblPr = table._tbl.tblPr # เข้าถึง properties ของตาราง
    tblCellMar = tblPr.xpath("w:tblCellMar") # หา tblCellMar element
    if tblCellMar: # ถ้าพบ
        tblCellMar = tblCellMar[0] # ใช้ element ที่พบ
    else: # ถ้าไม่พบ
        tblCellMar = OxmlElement('w:tblCellMar') # สร้าง tblCellMar element ใหม่
        tblPr.append(tblCellMar) # เพิ่มเข้าไปใน properties ของตาราง

    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)): # วนลูปในแต่ละด้าน
        elem = tblCellMar.find(qn(f'w:{side}')) # หา element สำหรับด้านนั้น
        if elem is None: # ถ้าไม่พบ
            elem = OxmlElement(f'w:{side}') # สร้าง element ใหม่
            tblCellMar.append(elem) # เพิ่มเข้าไป
        elem.set(qn('w:w'), str(val)) # กำหนดความกว้าง
        elem.set(qn('w:type'), 'dxa') # กำหนดประเภทหน่วยวัด

FIXED_METADATA = { # ข้อมูล metadata ที่ถูกกำหนดไว้ตายตัว
    "department": "ฝ่ายตรวจสอบธุรกิจเครือข่าย",
    "audit_task": "การบริหารจัดการด้านการให้บริการแก่ลูกค้าอย่างเป็นธรรม (Market Conduct)",
    "objective": "เพื่อให้มั่นใจว่าการบริหารจัดการเกี่ยวกับการให้บริการแก่ลูกค้ามีกระบวนการ การควบคุม\nและตรวจสอบการปฏิบัติงานที่คำนึงถึงการให้บริการอย่างเป็นธรรม",
    "control_code": "CT020400: การสอบทานกระบวนการขาย ตามหลักเกณฑ์ Market Conduct, Preventive Control",
}

def build_run_metadata(product: str, criteria_dir: str, output_file_path: str, voice_file: str | None) -> dict:
    """
    สร้าง metadata สำหรับตารางบนสุดของรายงาน
    product: ชื่อผลิตภัณฑ์
    criteria_dir: Directory ของไฟล์เกณฑ์
    output_file_path: Path ของไฟล์ output
    voice_file: Path ของไฟล์เสียง (ถ้ามี)
    return: dictionary ของ metadata
    """
    file_name = os.path.splitext(os.path.basename(output_file_path))[0] # ชื่อไฟล์ output (ไม่มีนามสกุล)

    criteria_file = f"{product}.json" # ชื่อไฟล์เกณฑ์
    keyword_file = f"{product}.csv" # ชื่อไฟล์ keyword
    voice = os.path.basename(voice_file) if voice_file else "" # ชื่อไฟล์เสียง (ถ้ามี)

    used_data = [] # list สำหรับเก็บข้อมูลที่ใช้
    if keyword_file: used_data.append(f"keyword: {keyword_file}") # เพิ่ม keyword file ถ้ามี
    if criteria_file: used_data.append(f"criteria: {criteria_file}") # เพิ่ม criteria file ถ้ามี
    if voice: used_data.append(f"voice: {voice}") # เพิ่ม voice file ถ้ามี

    run_date = datetime.now().strftime("%d/%m/%Y") # วันที่รันปัจจุบัน

    return { # คืน dictionary ของ metadata
        "department": FIXED_METADATA["department"],
        "audit_task": FIXED_METADATA["audit_task"],
        "objective": FIXED_METADATA["objective"],
        "control_code": FIXED_METADATA["control_code"],
        "file_name": file_name,
        "data_used": "\n".join(used_data), # เปลี่ยนเป็น join ด้วย newline
        "run_date": run_date,
    }

def insert_top_table_in_body(doc, metadata: dict):
    """
    แทรก 'ตารางข้อมูล' เป็นบรรทัดแรกของหน้าแรก (อยู่ในเนื้อเอกสาร ไม่ใช่ header)
    doc: Document object
    metadata: dictionary ของ metadata สำหรับตาราง
    """
    sec = doc.sections[0] # เข้าถึง section แรกของเอกสาร

    # คำนวณความกว้างที่พิมพ์ได้
    EMU_PER_INCH = 914400 # ค่าคงที่สำหรับแปลง EMU เป็นนิ้ว
    printable_emu = int(sec.page_width) - int(sec.left_margin) - int(sec.right_margin) # ความกว้างที่พิมพ์ได้ (EMU)
    printable_in = printable_emu / EMU_PER_INCH # ความกว้างที่พิมพ์ได้ (นิ้ว)

    left_w_in = 1.6 # ความกว้างคอลัมน์ซ้ายที่ต้องการ
    if printable_in < left_w_in + 1.0: # ถ้าความกว้างที่พิมพ์ได้น้อยเกินไป
        printable_in = left_w_in + 1.0 # ปรับความกว้างที่พิมพ์ได้
    right_w_in = printable_in - left_w_in # ความกว้างคอลัมน์ขวา

    # สร้างตาราง
    tbl = doc.add_table(rows=0, cols=2) # เพิ่มตารางใหม่
    try:
        tbl.style = "Table Grid" # กำหนด style ของตาราง
    except Exception:
        pass
    tbl.autofit = False # ปิด autofit
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT # จัดแนวตารางไปทางซ้าย

    # ข้อมูลแถวต่างๆ
    rows = [ # ข้อมูลสำหรับแต่ละแถวในตาราง
        ("หน่วยงาน", metadata.get("department", "")),
        ("งานตรวจสอบ", metadata.get("audit_task", "")),
        ("File Name", metadata.get("file_name", "")),
        ("Objective", metadata.get("objective", "")),
        ("Control Code", metadata.get("control_code", "")),
        ("ข้อมูลที่ใช้", metadata.get("data_used", "")),
        ("Run Date", metadata.get("run_date", "")),
    ]

    for k, v in rows: # วนลูปในแต่ละแถว
        cells = tbl.add_row().cells # เพิ่มแถวใหม่และเข้าถึง cells
        cells[0].width = Inches(left_w_in) # กำหนดความกว้างคอลัมน์ซ้าย
        cells[1].width = Inches(right_w_in) # กำหนดความกว้างคอลัมน์ขวา
        _set_cell(cells[0], k, bold=True) # ตั้งค่า cell ซ้าย (ตัวหนา)
        _set_cell(cells[1], v, bold=False) # ตั้งค่า cell ขวา

    # ย้ายตารางไปเป็น node ตัวแรกใน <w:body>
    body = doc._body._element # เข้าถึง body element
    el = tbl._element # เข้าถึง table element
    body.remove(el) # ลบ table element ออกจากตำแหน่งเดิม
    body.insert(0, el) # แทรก table element ที่ตำแหน่งแรกใน body

    # แทรกบรรทัดว่างถัดจากตารางเล็กน้อยให้ไม่ติดกับข้อความเดิม
    p = doc.add_paragraph() # เพิ่ม paragraph ว่าง
    p_el = p._element # เข้าถึง paragraph element
    body.remove(p_el) # ลบ paragraph element ออกจากตำแหน่งเดิม
    body.insert(1, p_el) # แทรก paragraph element ที่ตำแหน่งที่สองใน body


# ---------- render that APPENDS in LANDSCAPE and fits table ----------

def render_report_to_docx(
    original_path: str,
    report: Dict[str, Any],
    per_step_scores: Dict[int, Tuple[int, int]],
    output_file_path: str,
    output_dir: str,
    product_name: str = None  # เพิ่ม parameter # ชื่อผลิตภัณฑ์
):
    """
    สร้างรายงานผลการประเมินในรูปแบบ .docx โดยเพิ่มเนื้อหาในหน้าใหม่แบบแนวนอน
    original_path: Path ของไฟล์ .docx ต้นฉบับ
    report: รายงานผลลัพธ์จาก AI
    per_step_scores: คะแนนรายขั้นตอน
    output_file_path: Path ของไฟล์ .docx ที่จะบันทึก
    output_dir: Directory สำหรับบันทึกไฟล์ output
    product_name: ชื่อผลิตภัณฑ์ (สำหรับยกเว้นขั้นตอนที่ไม่คิดคะแนน)
    """
    # 1) เปิดไฟล์ input + สไตล์เริ่มต้น
    doc = Document(original_path) # เปิดเอกสาร .docx ต้นฉบับ
    _ensure_tahoma_styles(doc, body_size=11, h1_size=16) # ตั้งค่าฟอนต์เริ่มต้น

    # *** เพิ่มส่วนนี้: สร้างและแทรกตารางบนสุด ***
    # สร้าง metadata สำหรับตารางบน
    meta = build_run_metadata(
        product=product_name,  # ใช้ product_name จาก argument
        criteria_dir="criteria",
        output_file_path=output_file_path,
        voice_file=None
    )

    # รวมข้อมูล metadata จาก report (ถ้ามี)
    _meta_from_report = report.get("metadata", {}) if isinstance(report.get("metadata"), dict) else {} # ดึง metadata จากรายงาน

    top_table_meta = { # รวม metadata สำหรับตารางบนสุด
        "department": meta.get("department") or _meta_from_report.get("department") or "ฝ่ายตรวจสอบธุรกิจเครือข่าย",
        "audit_task": meta.get("audit_task") or _meta_from_report.get("audit_task") or "การบริหารจัดการด้านการให้บริการแก่ลูกค้าอย่างเป็นธรรม (Market Conduct)",
        "file_name": meta.get("file_name") or os.path.basename(original_path),
        "objective": meta.get("objective") or _meta_from_report.get("objective") or "เพื่อให้มั่นใจว่าการบริหารจัดการเกี่ยวกับการให้บริการแก่ลูกค้ามีกระบวนการ การควบคุมและตรวจสอบการปฏิบัติงานที่คำนึงถึงการให้บริการอย่างเป็นธรรม",
        "control_code": meta.get("control_code") or _meta_from_report.get("control_code") or "",
        "data_used": meta.get("data_used") or _meta_from_report.get("data_used") or "",
        "run_date": meta.get("run_date") or _meta_from_report.get("run_date") or time.strftime("%Y-%m-%d %H:%M"),
    }

    insert_top_table_in_body(doc, top_table_meta) # แทรกตารางบนสุดในเอกสาร
    # *** จบการเพิ่มตารางบนสุด ***

    # 2) เพิ่ม SECTION ใหม่ (ขึ้นหน้าใหม่) และตั้งเป็น LANDSCAPE (A4)
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE) # เพิ่ม section ใหม่
    new_sec.orientation = WD_ORIENT.LANDSCAPE # ตั้งค่าเป็นแนวนอน
    new_sec.page_width  = Mm(297)   # A4 แนวนอน # กำหนดความกว้างหน้ากระดาษ
    new_sec.page_height = Mm(210) # กำหนดความสูงหน้ากระดาษ
    new_sec.left_margin   = Inches(0.6) # กำหนดระยะขอบซ้าย
    new_sec.right_margin  = Inches(0.6) # กำหนดระยะขอบขวา
    new_sec.top_margin    = Inches(0.6) # กำหนดระยะขอบบน
    new_sec.bottom_margin = Inches(0.6) # กำหนดระยะขอบล่าง

    # 3) หัวเรื่อง
    model_name = report.get('_model_used', 'N/A') # ชื่อโมเดล AI ที่ใช้
    p_heading = doc.add_paragraph() # เพิ่ม paragraph สำหรับหัวเรื่อง
    r_heading = p_heading.add_run(f"ผลการประเมิน QA โดย AI — Model: {model_name}") # เพิ่มข้อความหัวเรื่อง
    _apply_run_font(r_heading, size_pt=18, bold=True) # ตั้งค่าฟอนต์หัวเรื่อง

    # 4) ตารางรายละเอียดรายขั้นตอน
    steps = report.get("steps", []) # ดึง steps จากรายงาน
    if not steps: # ถ้าไม่มี steps
        doc.add_paragraph("หมายเหตุ: ไม่พบข้อมูลรายละเอียดขั้นตอนจากโมเดล (steps ว่าง)") # เพิ่มข้อความแจ้งเตือน
    else:
        for step in steps: # วนลูปในแต่ละขั้นตอน
            doc.add_paragraph() # เพิ่ม paragraph ว่าง
            step_no = int(step.get("step", 0)) # เลขที่ขั้นตอน
            title = step.get("title", f"ขั้นตอน {step_no}") # title ของขั้นตอน
            score = step.get("score", "0/0") # คะแนนของขั้นตอน

            p = doc.add_paragraph() # เพิ่ม paragraph
            r = p.add_run(f"ขั้นตอน {step_no}) {title} — คะแนน: {score}") # เพิ่มข้อความขั้นตอนและคะแนน
            r.font.name = "Tahoma"; r.font.size = Pt(11); r.bold = True # ตั้งค่าฟอนต์

            _, expected = per_step_scores.get(step_no, (0, 0)) # จำนวนเกณฑ์ย่อยทั้งหมด
            doc.add_paragraph(f"เกณฑ์ย่อยทั้งหมด: {expected} ข้อ") # เพิ่มข้อความจำนวนเกณฑ์ย่อย

            tbl = doc.add_table(rows=1, cols=6) # เพิ่มตารางสำหรับรายละเอียดขั้นตอน
            tbl.style = "Table Grid" # กำหนด style
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER # จัดแนวตาราง
            tbl.autofit = False # ปิด autofit

            desired = [1.4, 0.75, 1.4, 1.6, 4.1, 0.75] # ความกว้างคอลัมน์ที่ต้องการ
            _apply_widths_fit(tbl, desired, new_sec, safety=0.92) # ปรับความกว้างตารางให้พอดีหน้า

            hdr = tbl.rows[0].cells # แถว header ของตาราง
            _set_cell_text(hdr[0], "Criterion", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell header
            _set_cell_text(hdr[1], "Status",   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(hdr[2], "Reason",   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(hdr[3], "Exact",    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(hdr[4], "Sentence", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(hdr[5], "Offset",   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            items = step.get("items", []) # ดึง items ย่อย
            if not isinstance(items, list): # ถ้าไม่ใช่ list
                items = [] # กำหนดเป็น list ว่าง

            if not items: # ถ้าไม่มี items
                row = tbl.add_row().cells # เพิ่มแถวใหม่
                _set_cell_text(row[0], "(ไม่มีข้อมูล)", align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell
                for j in range(1, 6):
                    _set_cell_text(row[j], "-", align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                for it in items: # วนลูปในแต่ละ item
                    row = tbl.add_row().cells # เพิ่มแถวใหม่
                    ev = it.get("evidence", {}) or {} # ดึง evidence
                    _set_cell_text(row[0], it.get("name","")) # ตั้งค่า cell
                    _set_cell_text(row[1], it.get("status",""), align=WD_ALIGN_PARAGRAPH.CENTER)
                    _set_cell_text(row[2], it.get("reason",""))
                    _set_cell_text(row[3], ev.get("exact",""))
                    _set_cell_text(row[4], ev.get("sentence",""))
                    off = ev.get("offset", []) # ดึง offset
                    _set_cell_text(
                        row[5],
                        ", ".join(map(str, off)) if isinstance(off, list) else str(off),
                        align=WD_ALIGN_PARAGRAPH.CENTER
                    )

            for row in tbl.rows[1:]: # วนลูปในแถวข้อมูล (ไม่รวม header)
                _shrink_font(row.cells[2], pt=10) # ลดขนาดฟอนต์
                _shrink_font(row.cells[3], pt=10)
                _shrink_font(row.cells[4], pt=10)

    # 5) สรุปรวม
    doc.add_paragraph() # เพิ่ม paragraph ว่าง
    p = doc.add_paragraph("สรุปรวม") # เพิ่ม paragraph "สรุปรวม"
    p.style = "Heading 1" # กำหนด style เป็น Heading 1

    for r in p.runs: # วนลูปใน run ของ paragraph
        r.font.name = "Tahoma" # กำหนดชื่อฟอนต์
        r.font.size = Pt(16) # กำหนดขนาดฟอนต์
        r.bold = True # กำหนดตัวหนา

    pf = p.paragraph_format # เข้าถึง paragraph format
    pf.space_before = Pt(12) # กำหนดระยะห่างก่อน paragraph
    pf.space_after  = Pt(12) # กำหนดระยะห่างหลัง paragraph

    sumsec = report.get("summary", {}) or {} # ดึง summary จากรายงาน
    tb = doc.add_table(rows=1, cols=3) # เพิ่มตารางสำหรับสรุปคะแนน
    tb.style = "Table Grid" # กำหนด style
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER # จัดแนวตาราง
    tb.autofit = False # ปิด autofit
    _apply_widths_fit(tb, [1.00, 4.00, 1.50], new_sec, safety=0.88) # ปรับความกว้างตาราง
    _set_table_cell_margins(tb, left=40, right=40, top=20, bottom=20) # ตั้งค่าระยะขอบ cell

    h = tb.rows[0].cells # แถว header ของตาราง
    _set_cell_text(h[0], "Step", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell header
    _set_cell_text(h[1], "ชื่อขั้นตอน", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(h[2], "คะแนน", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ในส่วนของการสร้างตารางสรุปคะแนนรายขั้นตอน (ประมาณบรรทัด 580-600)
    # แก้ไขการแสดงผลคะแนนในตาราง
    excluded_steps = EXCLUDED_STEPS_BY_PRODUCT.get(product_name, set()) # ดึงขั้นตอนที่ถูกยกเว้น

    for s in steps: # วนลูปในแต่ละขั้นตอน
        r = tb.add_row().cells # เพิ่มแถวใหม่
        step_no = int(s.get("step", 0)) # เลขที่ขั้นตอน

        _set_cell_text(r[0], s.get("step",""), align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell
        _set_cell_text(r[1], s.get("title",""))

        # แสดงคะแนนตามที่กำหนด (รวมหมายเหตุสำหรับขั้นตอนที่ไม่คิด)
        if step_no in excluded_steps: # ถ้าขั้นตอนนี้ถูกยกเว้น
            _set_cell_text(r[2], "0/0 (ไม่คิดคะแนน)", align=WD_ALIGN_PARAGRAPH.CENTER) # แสดงคะแนนเป็น "0/0 (ไม่คิดคะแนน)"
        else:
            _set_cell_text(r[2], s.get("score",""), align=WD_ALIGN_PARAGRAPH.CENTER) # แสดงคะแนนปกติ

    doc.add_paragraph() # เพิ่ม paragraph ว่าง

    p_total = doc.add_paragraph(f"คะแนนรวม: {sumsec.get('_score_total', '0/0')}") # เพิ่ม paragraph คะแนนรวม
    p_total.paragraph_format.space_before = Pt(0) # กำหนดระยะห่างก่อน paragraph
    p_total.paragraph_format.space_after  = Pt(6) # กำหนดระยะห่างหลัง paragraph

    p_pct = doc.add_paragraph(f"คิดเป็นเปอร์เซ็นต์ (%): {sumsec.get('_score_percent', '0%')}") # เพิ่ม paragraph เปอร์เซ็นต์

    level = classify_overall(report) # จัดระดับการประเมินโดยรวม
    pp = doc.add_paragraph(); rr = pp.add_run(f"การประเมินโดยรวม: {level}") # เพิ่ม paragraph การประเมินโดยรวม
    rr.font.name = "Tahoma"; rr.font.size = Pt(11); rr.bold = True # ตั้งค่าฟอนต์

    reason = (sumsec.get("narrative_reason") or "").strip() # เหตุผลการจัดระดับ
    if reason: # ถ้ามีเหตุผล
        p = doc.add_paragraph("เหตุผลการจัดระดับ") # เพิ่ม paragraph "เหตุผลการจัดระดับ"
        for r in p.runs:
            r.font.name = "Tahoma"; r.font.size = Pt(12); r.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        p_body = doc.add_paragraph(reason) # เพิ่ม paragraph เนื้อหาเหตุผล
        for r in p_body.runs:
            r.font.name = "Tahoma"; r.font.size = Pt(11)

    # จุดเด่น
    strengths = sumsec.get("strengths") or [] # จุดเด่น
    if strengths: # ถ้ามีจุดเด่น
        p = doc.add_paragraph("จุดเด่น") # เพิ่ม paragraph "จุดเด่น"
        for r in p.runs:
            r.font.name = "Tahoma"; r.font.size = Pt(12); r.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        for s in strengths: # วนลูปในจุดเด่นแต่ละรายการ
            p_b = doc.add_paragraph(s, style="List Bullet") # เพิ่ม paragraph แบบ bullet
            for r in p_b.runs: r.font.name = "Tahoma"; r.font.size = Pt(11)

    # ควรปรับปรุง
    improvements = sumsec.get("improvements") or [] # สิ่งที่ควรปรับปรุง
    if improvements: # ถ้ามีสิ่งที่ควรปรับปรุง
        p = doc.add_paragraph("ควรปรับปรุง") # เพิ่ม paragraph "ควรปรับปรุง"
        for r in p.runs:
            r.font.name = "Tahoma"; r.font.size = Pt(12); r.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

        for s in improvements: # วนลูปในสิ่งที่ควรปรับปรุงแต่ละรายการ
            p_b = doc.add_paragraph(s, style="List Bullet") # เพิ่ม paragraph แบบ bullet
            for r in p_b.runs: r.font.name = "Tahoma"; r.font.size = Pt(11)

    doc.add_paragraph() # เพิ่ม paragraph ว่าง
    p_note = doc.add_paragraph() # เพิ่ม paragraph สำหรับหมายเหตุ
    run_note = p_note.add_run("หมายเหตุ: ตัวอย่างการตัดสินใจของ AI") # เพิ่มข้อความหมายเหตุ
    run_note.bold = True # ตัวหนา
    run_note.font.name = "Tahoma" # ชื่อฟอนต์
    run_note.font.size = Pt(11) # ขนาดฟอนต์

    # ตามด้วย bullet/บรรทัด "ตัวอย่างการตัดสินใจ..." ของคุณเหมือนเดิม
    for line in DECISION_GUIDE_TEXT: # วนลูปในข้อความแนะนำการตัดสินใจ
        doc.add_paragraph(f"• {line}") # เพิ่ม paragraph แบบ bullet

    # *** เพิ่มส่วน Signature Table ที่ด้านล่างสุด ***
    # เพิ่มพื้นที่ว่าง 4-5 บรรทัดก่อนส่วนลายเซ็น
    for _ in range(4): # เพิ่ม paragraph ว่าง 4 ครั้ง
        doc.add_paragraph()

    # สร้างตารางสำหรับลายเซ็น (2 คอลัมน์)
    signature_table = doc.add_table(rows=3, cols=2) # เพิ่มตารางลายเซ็น
    signature_table.style = "Table Grid" # กำหนด style
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER # จัดแนวตาราง
    signature_table.autofit = False # ปิด autofit

    # กำหนดความกว้างของตาราง signature
    current_printable_in = (new_sec.page_width.inches - new_sec.left_margin.inches - new_sec.right_margin.inches) * 0.9 # ความกว้างที่พิมพ์ได้
    col_width = current_printable_in / 2 # ความกว้างแต่ละคอลัมน์
    _apply_widths_fit(signature_table, [col_width, col_width], new_sec, safety=0.9) # ปรับความกว้างตาราง

    # แถวที่ 1: จุดๆ สำหรับลายเซ็น
    cells = signature_table.rows[0].cells # เข้าถึง cells ในแถวแรก
    _set_cell(cells[0], "........................", align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell
    _set_cell(cells[1], "........................", align=WD_ALIGN_PARAGRAPH.CENTER)

    # แถวที่ 2: ชื่อตำแหน่ง
    cells = signature_table.rows[1].cells # เข้าถึง cells ในแถวที่สอง
    _set_cell(cells[0], "ผู้จัดเตรียมกระดาษทำการ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell
    _set_cell(cells[1], "ผู้สอบทานกระดาษทำการ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # แถวที่ 3: วงเล็บสำหรับวันที่
    cells = signature_table.rows[2].cells # เข้าถึง cells ในแถวที่สาม
    _set_cell(cells[0], "( ....../....../...... )", align=WD_ALIGN_PARAGRAPH.CENTER) # ตั้งค่า cell
    _set_cell(cells[1], "( ....../....../...... )", align=WD_ALIGN_PARAGRAPH.CENTER)

    # ปรับแต่งตาราง - ลบเส้นขอบ
    for row in signature_table.rows: # วนลูปในแต่ละแถว
        for cell in row.cells: # วนลูปในแต่ละ cell
            tc = cell._tc # เข้าถึง tc element
            tcPr = tc.get_or_add_tcPr() # เข้าถึงหรือเพิ่ม tcPr
            tcBorders = OxmlElement('w:tcBorders') # สร้าง tcBorders element
            for border_name in ['top', 'left', 'bottom', 'right']: # วนลูปในแต่ละเส้นขอบ
                border = OxmlElement(f'w:{border_name}') # สร้าง border element
                border.set(qn('w:val'), 'nil') # กำหนดค่าเป็น 'nil' (ไม่มีเส้นขอบ)
                tcBorders.append(border) # เพิ่ม border เข้าไป
            tcPr.append(tcBorders) # เพิ่ม tcBorders เข้าไปใน tcPr

    # 6) save แบบปลอดภัย
    os.makedirs(output_dir, exist_ok=True) # สร้าง directory output ถ้ายังไม่มี
    tmp = output_file_path + ".tmp" # สร้างชื่อไฟล์ชั่วคราว
    doc.save(tmp) # บันทึกเอกสารลงไฟล์ชั่วคราว
    os.replace(tmp, output_file_path) # แทนที่ไฟล์เดิมด้วยไฟล์ชั่วคราว

# --- Main Execution Logic ---

def main():
    parser = argparse.ArgumentParser(description="AI-Powered QA Summarizer for Call Transcripts") # สร้าง ArgumentParser
    parser.add_argument("--product", required=True, help="Product name matching the criteria JSON file (e.g., 'debit_card')") # เพิ่ม argument สำหรับชื่อผลิตภัณฑ์
    parser.add_argument("--criteria-dir", default="criteria", help="Directory containing criteria JSON files") # เพิ่ม argument สำหรับ directory เกณฑ์
    parser.add_argument("--input-dir", default="transcript_with_highlight", help="Directory containing input .docx files") # เพิ่ม argument สำหรับ directory input
    parser.add_argument("--output-dir", default="transcript_with_highlight_and_ai_summarize", help="Base directory for AI summary output") # เพิ่ม argument สำหรับ directory output
    parser.add_argument("--model", default="google/gemma-3-27b-it:free", help="The primary AI model to use") # เพิ่ม argument สำหรับโมเดล AI หลัก
    parser.add_argument("--fallback-model", default="google/gemma-3-12b-it:free", help="Fallback model if the primary fails") # เพิ่ม argument สำหรับโมเดล AI สำรอง
    args = parser.parse_args() # Parse arguments

    client = get_openai_client() # สร้าง OpenAI client
    models = [args.model, args.fallback_model] # list ของโมเดลที่จะใช้

    try:
        criteria_data = load_criteria_from_json(args.product, args.criteria_dir) # โหลดข้อมูลเกณฑ์
    except FileNotFoundError as e:
        logging.error(e) # บันทึก log เมื่อไม่พบไฟล์เกณฑ์
        return # ออกจากโปรแกรม

    input_path = os.path.join(os.path.dirname(__file__), args.input_dir, args.product) # สร้าง path สำหรับ input directory

    output_dir_ai = os.path.join(args.output_dir, f"{args.product}_final_output") # สร้าง path สำหรับ output directory ของ AI
    os.makedirs(output_dir_ai, exist_ok=True) # สร้าง output directory ถ้ายังไม่มี

    files = glob.glob(os.path.join(input_path, "*.docx")) # ค้นหาไฟล์ .docx ทั้งหมดใน input directory
    files = [f for f in files if not os.path.basename(f).startswith("~$")] # กรองไฟล์ชั่วคราวออก

    if not files: # ถ้าไม่พบไฟล์ .docx
        logging.warning(f"No .docx files found in '{input_path}'.") # บันทึก log แจ้งเตือน
        return # ออกจากโปรแกรม

    primary_model = args.model # โมเดลหลัก
    fallback_model = args.fallback_model # โมเดลสำรอง
    model_list = [m for m in [primary_model, fallback_model] if m] # list ของโมเดลที่ใช้งานได้

    steps = ["Load transcript", "Build prompt", "Call AI", "Score", "Render"] # ขั้นตอนการทำงาน (ไม่ได้ใช้โดยตรงใน tqdm)
    with logging_redirect_tqdm(): # redirect logging ให้แสดงผลเหนือ progress bar
        total = len(files) # จำนวนไฟล์ทั้งหมด
        for idx, docx_path in enumerate(files, start=1): # วนลูปในแต่ละไฟล์ .docx
            fname = os.path.basename(docx_path) # ชื่อไฟล์

            logging.info("\n" + "-"*100) # บันทึก log เส้นแบ่ง
            logging.info(f"▶ Processing file {idx}/{total}: {fname}") # บันทึก log การประมวลผลไฟล์
            logging.info("-"*100)
            logging.info("-"*100 + "\n")

            t0 = time.time() # บันทึกเวลาเริ่มต้น
            final_output_file_path = None # Path ของไฟล์ output สุดท้าย
            skip_current_file = False # flag สำหรับข้ามไฟล์ปัจจุบัน

            steps = ["Load transcript", "Process with AI", "Score", "Render"] # ขั้นตอนการทำงานสำหรับ progress bar
            with tqdm(total=len(steps), desc=f"[{idx}/{total}] {fname}", unit="step", leave=True) as bar: # สร้าง progress bar

                # 1) Load transcript
                transcript = read_docx(docx_path) # อ่าน transcript จากไฟล์ .docx
                if not transcript or not transcript.strip(): # ถ้า transcript ว่างเปล่า
                    logging.warning(f"Skipping empty transcript: {fname}") # บันทึก log แจ้งเตือน
                    skip_current_file = True # ตั้งค่าให้ข้ามไฟล์
                else:
                    logging.info(f"Loaded transcript: {len(transcript)} characters") # บันทึก log ความยาว transcript
                bar.update(1) # อัปเดต progress bar

                if not skip_current_file: # ถ้าไม่ข้ามไฟล์
                    # 2) Process with AI (ใช้ฟังก์ชันใหม่)
                    bar.set_postfix(status="Processing with AI...") # ตั้งค่าข้อความใน progress bar
                    report = process_single_file_with_retries(transcript, criteria_data, client, model_list) # ประมวลผลด้วย AI
                    bar.update(1) # อัปเดต progress bar

                    if not report: # ถ้าไม่ได้รับรายงานจาก AI
                        logging.error(f"All processing attempts failed for: {fname}") # บันทึก log error
                        skip_current_file = True # ตั้งค่าให้ข้ามไฟล์

                    if not skip_current_file: # ถ้าไม่ข้ามไฟล์
                        # Normalize report
                        report = normalize_report(report, criteria_data) # ปรับโครงสร้างรายงาน

                        # 3) Score
                        scored_report, per_step_scores = compute_scores(report, args.product) # คำนวณคะแนน
                        bar.update(1) # อัปเดต progress bar

                        # 4) Render
                        output_file_name = f"ai_summary_{fname}" # ชื่อไฟล์ output
                        final_output_file_path = os.path.join(output_dir_ai, output_file_name) # Path ของไฟล์ output
                        render_report_to_docx(docx_path, scored_report, per_step_scores, final_output_file_path, output_dir_ai, args.product) # สร้างรายงาน .docx
                        bar.update(1) # อัปเดต progress bar

                # Update progress bar if skipped
                if skip_current_file: # ถ้าข้ามไฟล์
                    if bar.n < bar.total: # ถ้า progress bar ยังไม่เต็ม
                        bar.update(bar.total - bar.n) # อัปเดต progress bar ให้เต็ม

            # Summary for this file
            if skip_current_file: # ถ้าข้ามไฟล์
                logging.info(f"⭕ Skipped: {fname}") # บันทึก log ว่าข้ามไฟล์
            else:
                elapsed = (time.time() - t0) / 60 # เวลาที่ใช้ในการประมวลผล
                final_output_file_path = final_output_file_path or "(no output path)" # กำหนด path output
                logging.info(f"✅ Completed: {fname} -> {final_output_file_path} | elapsed {elapsed:.2f} min") # บันทึก log ว่าประมวลผลสำเร็จ

            # *** เพิ่มการ delay ระหว่างไฟล์ที่นี่ ***
            if idx < total:  # ไม่ delay หลังไฟล์สุดท้าย
                logging.info(f"Waiting 5 seconds before processing next file...") # บันทึก log การรอ
                time.sleep(5) # รอ 5 วินาที

    logging.info(f"\n🎉 All files processed! Total files: {total}") # บันทึก log เมื่อประมวลผลไฟล์ทั้งหมดเสร็จสิ้น

if __name__ == "__main__":
    main() # เรียกฟังก์ชัน main เมื่อ script ถูกรันโดยตรง
