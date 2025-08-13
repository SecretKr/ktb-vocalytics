# ลบ import fuzzywuzzy ออก
import csv
import glob
import os
import re

import docx
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv() 

# Function to read text from .docx file
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to append text to an existing .docx file
def append_to_docx(file_path, content, model_name):
    doc = docx.Document(file_path)
    doc.add_paragraph(f"\n--- สรุปเนื้อหาโดย AI (เพิ่มเติม) โดยใช้โมเดล {model_name} ---\n")
    doc.add_paragraph(content)
    doc.save(file_path)

# Function to read keywords from a CSV file
def read_keywords_from_csv(csv_file_path):
    keywords = []
    try:
        with open(csv_file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if row:
                    keywords.extend([kw.strip() for kw in row if kw.strip()])
    except FileNotFoundError:
        print(f"Keyword file not found: {csv_file_path}")
    return list(set(keywords))

# แก้ไขฟังก์ชันนี้ให้ค้นหาแบบตรงตัวแทน
def add_keyword_highlighting(text, keywords):
    highlighted_text = text
    for keyword in keywords:
        # ใช้ re.sub เพื่อแทนที่คำที่ตรงกันแบบตรงตัว
        highlighted_text = re.sub(r'\b' + re.escape(keyword) + r'\b', r'**' + keyword + r'**', highlighted_text, flags=re.IGNORECASE)
    return highlighted_text

if __name__ == '__main__':
    openrouter_api_key = os.getenv("OPENROUTER_API_KEY")
    if not openrouter_api_key:
        raise ValueError("Please set the OPENROUTER_API_KEY environment variable.")

    openrouter_model = 'google/gemma-3-27b-it:free'
    
    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=openrouter_api_key,
    )

    script_dir = os.path.dirname(__file__)
    input_directory = os.path.join(script_dir, "transcript_with_highlight")
    keywords_directory = os.path.join(script_dir, os.pardir, "keywords")
    
    os.makedirs(input_directory, exist_ok=True)
    
    print(f"Attempting to find .docx files in: {os.path.abspath(input_directory)}")
    all_docx_files = glob.glob(os.path.join(input_directory, "*.docx"))
    docx_files = [f for f in all_docx_files if not os.path.basename(f).startswith('~$')]
    print(f"Files found: {docx_files}")

    # personal_loan_keywords = read_keywords_from_csv(os.path.join(keywords_directory, "personal_loan.csv"))
    # debit_card_keywords = read_keywords_from_csv(os.path.join(keywords_directory, "debit_card.csv"))
    # keywords_to_use = list(set(personal_loan_keywords + debit_card_keywords))

    debit_card_keywords = read_keywords_from_csv(os.path.join(keywords_directory, "debit_card.csv"))
    keywords_to_use = list(set(debit_card_keywords))

    if not docx_files:
        print(f"No .docx files found in {os.path.abspath(input_directory)}")
    else:
        for file_path in docx_files:
            file_name = os.path.basename(file_path)
            print(f"Processing {file_name} for summarization using model: {openrouter_model}...")
            
            try:
                document_text = read_docx(file_path)

                if not document_text.strip():
                    print(f"The document {file_name} is empty. Skipping summarization.")
                    continue
                else:
                    prompt_template = """
                คุณคือผู้เชี่ยวชาญด้านการวิเคราะห์บทสนทนา Call Center เพื่อวัตถุประสงค์ด้านการตรวจสอบคุณภาพงานขายทางโทรศัพท์  
                เนื้อหาที่จะวิเคราะห์เป็นข้อความที่แปลงจากเสียง (Voice to Text) ซึ่งอาจไม่สมบูรณ์หรือมีความผิดเพี้ยนของคำบางส่วน

                ### งานของคุณ:
                1. **สรุปเนื้อหาอย่างเป็นระบบ แยกตามหัวข้อ**
                2. **เน้นดึงรายละเอียดที่เป็นสาระสำคัญ แม้เนื้อหาจะมี noise**
                3. **ลำดับเหตุการณ์ชัดเจน**
                4. **ระบุผลิตภัณฑ์ที่ถูกเสนอ, เงื่อนไข, และข้อมูลสำคัญต่าง ๆ อย่างครบถ้วน**
                5. **หากพบสิ่งผิดปกติ เช่น คำพูดซ้ำซ้อน ผิดเพี้ยน หรือขัดแย้ง ให้ระบุว่า "ไม่ชัดเจน" แต่ไม่ควรสรุปว่า "ไม่มีข้อมูล" โดยเด็ดขาด หากยังมีบางส่วนพอวิเคราะห์ได้**

                ---

                ### โปรดสรุปเนื้อหาดังนี้:

                1. **ลำดับเหตุการณ์ (Timeline Summary)**
                   - สรุปเหตุการณ์สำคัญแบบเรียงตามลำดับเวลา (เช่น เริ่มจากทักทาย → แนะนำผลิตภัณฑ์ → ตอบคำถามลูกค้า ฯลฯ)
                   - หากพบการวนลูป พูดซ้ำ หรือหลุดบริบท ให้ระบุ [เนื้อหาซ้ำ/ไม่ต่อเนื่อง]

                2. **การทักทายและเริ่มต้นบทสนทนา**
                   - มีการทักทายอย่างไร
                   - พนักงานเปิดบทสนทนาแบบสุภาพหรือไม่

                3. **รายละเอียดผลิตภัณฑ์ที่เสนอขาย**
                   - ประเภทผลิตภัณฑ์ (บัตรเดบิต, สินเชื่อ, บริการฝากเงิน ฯลฯ)
                   - จุดเด่น/สิทธิประโยชน์/ข้อเสนอพิเศษที่ถูกกล่าวถึง
                   - เงื่อนไข เช่น อัตราดอกเบี้ย ระยะเวลาฝาก ข้อจำกัดการถอน ฯลฯ
                   - เนื้อหาด้านประกันภัย (หากเกี่ยวข้อง)

                4. **การให้คำแนะนำลูกค้า (Responsible Lending / Compliance)**
                   - พนักงานมีการเตือนความเสี่ยง หรือให้ข้อมูลด้านความรับผิดชอบทางการเงินหรือไม่
                   - การระบุเงื่อนไขหรือข้อควรระวัง เช่น การถอนก่อนกำหนด, ข้อจำกัดด้านสิทธิประโยชน์

                5. **การตอบโต้ของลูกค้า**
                   - ลูกค้าแสดงความสนใจ, ถามกลับ, ปฏิเสธ, หรือไม่ตอบอย่างไร
                   - คำพูดสำคัญของลูกค้า (ระบุได้เท่าที่ปรากฏ)

                6. **การปิดการขาย / สรุปท้ายบทสนทนา**
                   - พนักงานมีความพยายามสรุป/ติดตาม/เสนอทางเลือกต่อหรือไม่
                   - มีการกล่าวขอบคุณหรือปิดท้ายอย่างสุภาพหรือไม่

                7. **สิ่งที่ควรสังเกตเพิ่มเติม**
                   - ความผิดเพี้ยนทางภาษาที่อาจเกิดจาก Voice to Text
                   - คำที่ดูเหมือนพูดซ้ำ/ผิดบริบท
                   - ข้อความหลุดประเด็นหรือไม่เกี่ยวข้อง

                ---

                **ข้อกำหนดเพิ่มเติม:**
                - เขียนเป็นภาษาไทยทางการ เหมาะสำหรับใช้ในรายงาน
                - หากบางส่วน [ฟังไม่ชัด], [ขาดหาย], หรือ [ไม่สมเหตุสมผล] ให้ใส่ระบุไว้ตรงนั้น
                - อย่าตัดทอนเนื้อหาโดยรวม แต่ให้ดึงแก่นสำคัญที่วิเคราะห์ได้

                ---

                บทสนทนา:
                {document_text}
                """

                    prompt = prompt_template.format(document_text=document_text)
                    print(f"Sending text from {file_name} to AI for summarization...")
                    
                    response = client.chat.completions.create(
                        model=openrouter_model,
                        messages=[{'role': 'user', 'content': prompt}],
                        temperature=0.7,
                    )
                    
                    summarized_content = response.choices[0].message.content

                    print(f"\n--- สรุปเนื้อหาโดย AI สำหรับ {file_name} (Model: {openrouter_model}) ---\n")
                    print(summarized_content.strip())
                    print("\n" + "-"*30)

                    # เรียกใช้ฟังก์ชันที่แก้ไขแล้ว
                    highlighted_content = add_keyword_highlighting(summarized_content, keywords_to_use)

                    append_to_docx(file_path, highlighted_content, openrouter_model)
                    print(f"Summarized and highlighted content appended to {file_path}")

            except FileNotFoundError:
                print(f"Error: File not found at {file_path}")
            except Exception as e:
                print(f"An error occurred while processing {file_name}: {e}")